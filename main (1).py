from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, BackgroundTasks, Request
from sqlalchemy.orm import Session
from sql_app import crud, models, schemas
from langchain_community.llms import Ollama
from sql_app.database import SessionLocal, engine
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings.fastembed import FastEmbedEmbeddings
from langchain_community.vectorstores import Chroma
from sqlalchemy.orm import Session
from sql_app import crud, models, schemas
from sql_app.database import SessionLocal, engine
import os
import asyncio
from uuid import uuid4
from fastapi.responses import JSONResponse
from langchain.prompts import PromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from fastapi.middleware.cors import CORSMiddleware
import json
import ollama
from fastapi.responses import StreamingResponse
from fastapi.responses import FileResponse

from typing import AsyncGenerator
from fastapi.responses import HTMLResponse
import pypdf2htmlEX
import pdfplumber
import io
import docx
import re
import subprocess
import traceback
import logging
import pytz
from datetime import datetime, timezone, timedelta
import docx2txt
import shutil
from pydantic import BaseModel


def get_logger(log_file="logs.log"):
    """
    Creates and returns a logger that writes logs to a specified file with time in IST format.

    Parameters:
        log_file (str): The name of the log file. Defaults to 'logs'.

    Returns:
        logging.Logger: Configured logger instance.
    """
    logger = logging.getLogger(log_file)
    logger.setLevel(logging.DEBUG)

    # Ensure we don't add multiple handlers if the logger already exists        # File handler for logging
    file_handler = logging.FileHandler(log_file)
    file_handler.setLevel(logging.DEBUG)

        # Formatter for log messages with IST time
    formatter = logging.Formatter(
            "%(asctime)s - %(levelname)s - %(message)s"
        )

        # Custom time format function to convert to IST
    def utc_to_ist(*args):
        utc_dt = datetime.now(timezone.utc)  # UTC time with timezone info
        ist_dt = utc_dt + timedelta(hours=5, minutes=30)  # Convert to IST
        return ist_dt.timetuple()

    logging.Formatter.converter = utc_to_ist

    file_handler.setFormatter(formatter)

        # Add the handler to the logger
    logger.addHandler(file_handler)
    return logger

models.Base.metadata.create_all(bind=engine)

app = FastAPI()

# Configure CORS settings
origins = [
    "http://localhost",  # Frontend running on the same machine
    "http://localhost:3000",  # React or other frontend frameworks running on different ports
    "http://example.com",
    "*",# Your production domain
    # Add more origins as needed
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,  # List of allowed origins
    allow_credentials=True,  # To allow cookies or authentication headers
    allow_methods=["*"],  # List of allowed HTTP methods (GET, POST, etc.), "*" means all methods
    allow_headers=["*"],  # List of allowed headers, "*" means all headers
)

cached_llm = Ollama(model="llama3.1:70b")

# embedding_model_name = "all-MiniLM-L6-v2"  # Replace with your actual model name
embedding = FastEmbedEmbeddings()


# embedding_model = Ollama(model="llama3.2")  # Adjust the model name as necessary

# # Check if embedding model is initialized
# if embedding._model is None:
#     raise ValueError("Embedding model is not initialized properly.")

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1024, chunk_overlap=80, length_function=len, is_separator_regex=False
)

raw_prompt = PromptTemplate.from_template(
    """ 
    <s>[INST] You are a technical assistant good at searching docuemnts. If you do not have an answer from the provided information say so. [/INST] </s>
    [INST] {input}
           Context: {context}
           Answer:
    [/INST]
"""
)

folder_path = "db"


# Directory to store uploaded PDFs
UPLOAD_DIR = "pdf"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Dependency to get a database session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()
        

# # Assuming cached_llm.invoke is a generator that yields response chunks
# async def stream_response(query: str) -> AsyncGenerator[str, None]:
#     print(f"query: {query}")
    
#     # Simulate the streaming response from the LLM (replace this with your actual streaming logic)
#     async for chunk in cached_llm.invoke(query):
#         yield chunk
#         await asyncio.sleep(0.1)  

@app.post("/ask_pdf")
async def ask_pdf_post(request: schemas.AskQueryRequest, db: Session = Depends(get_db)):
    uuid = request.uuid
    query = request.query
    logger=get_logger()
    logger.info(f'{query}{uuid}')
    print(f"Post /ask_pdf called with UUID: {uuid} and query: {query}")
    
    # Step 1: Check if the UUID exists in the database
    pdf_entry = db.query(models.PDF).filter(models.PDF.source_id == str(uuid)).first()

    if not pdf_entry:
        # UUID not found in the database, return an error message
        raise HTTPException(status_code=404, detail=f"UUID {uuid} is not associated with any document.")

    # Step 2: Extract chunks from the data column
    response_data = pdf_entry.data  # Assuming `data` column holds JSON response
    if "chunks" not in response_data:
        k = 20
        # raise HTTPException(status_code=400, detail=f"Chunks information missing for UUID {uuid}.")

    k = min(response_data["chunks"],100)  # Dynamically set k based on chunks
    print(f"Number of chunks (k): {k}")

    print("Loading vector store")
    vector_store = Chroma(persist_directory=folder_path, 
                          embedding_function=embedding,
                          collection_name= str(uuid)
                          )
    
    print(vector_store)

    print("Filtering documents by UUID")
    search_results = vector_store.similarity_search(
        query=query,
        k = k,
        filter={"uuid": str(uuid)}
    )

    filtered_results = [doc for doc in search_results if doc.metadata.get("uuid") == str(uuid)]

    # Check if any documents were found for the given UUID
    if not filtered_results:
        raise HTTPException(status_code=404, detail=f"No documents found for UUID: {uuid}")

    print(f"Found {len(filtered_results)} documents for UUID: {uuid}")


    print("Creating retriever and document chain")
    retriever = vector_store.as_retriever(
        search_type="similarity_score_threshold",
        search_kwargs={
            "k": 2,  # Number of results to return
            "score_threshold": 0.2,  # Similarity score threshold
        },
    )

    document_chain = create_stuff_documents_chain(cached_llm, raw_prompt)
    chain = create_retrieval_chain(retriever, document_chain)

    # Perform the search query on the filtered document chunks
    result = chain.invoke({
        "input": query,
        "documents": filtered_results  # Pass the filtered documents to the chain
    })

    
    # Extracting the sources
    sources = []
    for doc in result["context"]:
        sources.append(
            {"page_number": doc.metadata["page"], "file_name": doc.metadata["file_path"], "page_content": doc.page_content}
        )

    # Preparing the response
    response_answer = {"answer": result["answer"], "sources": sources}
    file_handler.close()
    logger.removeHandler(file_handler)
    return response_answer

      
        
@app.post("/pdf")
async def pdf_post(file: UploadFile = File(...), db: Session = Depends(get_db)):
    file_name = file.filename
    save_file = os.path.join("pdf", file_name)
    
    # Save the uploaded file
    with open(save_file, "wb") as f:
        f.write(await file.read())
    
    print(f"filename: {file_name}")

    # Load and split the PDF
    loader = PDFPlumberLoader(save_file)
    docs = loader.load_and_split()
    if not docs:
        return HTTPException(status_code=502, detail="something wrong with the pdf")
    print(f"docs len={len(docs)}")

    chunks = text_splitter.split_documents(docs)
    print(f"chunks len={len(chunks)}")
    

    # Generate a unique ID (UUID)
    unique_id = str(uuid4())
    
    for chunk in chunks:
        chunk.metadata['uuid'] = unique_id
    
        # Store the embeddings with the unique ID
    vector_store = Chroma.from_documents(
        documents=chunks,
        embedding=embedding,
        persist_directory=folder_path,
        collection_name= str(unique_id)
    )
    print(vector_store)
    
    vector_store.persist()

    # Prepare the response data
    response = {
        "status": "Successfully Uploaded",
        "filename": file_name,
        "uuid": unique_id,
        "doc_len": len(docs),
        "chunks": len(chunks),
    }

    # Create PDF entry in the database
    pdf_data = schemas.PDFCreate(
        pdf_name=file_name, 
        pdf_data=str(docs), 
        source_id=unique_id,
        data=response  # Store the response in the "data" column
    )
    db_pdf = crud.create_pdf(db=db, pdf=pdf_data)

    response["pdf_id"] = db_pdf.id  # Add PDF ID to response

    return JSONResponse(content=response)


@app.post("/ai")
async def ai(query: str, db: Session = Depends(get_db)):
    print("hello")
    print("Post /ai called")
    # json_content = request.json
    # query = json_content.get("query")

    print(f"query: {query}")

    response = cached_llm.invoke(query)

    print(response)

    response_answer = {"answer": response}
    return response_answer



@app.post("/resume")
async def parse_resume(file: UploadFile = File(...)):
    try:
        logger=get_logger()
        # Check file type
        logger.info(f'The filename is {file.filename}')
        print('The file name is',file.filename)
        print('The file name is',file.filename)
        print('The file name is',file.filename)

        if not file.filename.endswith(('.pdf', '.docx','.doc')):
            raise HTTPException(status_code=400, detail="Unsupported file format")

        # Read the file content
        if file.filename.endswith('.doc'):
            try:
                with open(f"{file.filename}", "wb") as buffer:
                    shutil.copyfileobj(file.file, buffer)
            except Exception as e:
                logger.exception()
                return JSONResponse(content={"message": f"Error: {str(e)}"}, status_code=400)
        
        file_content = await file.read()

        # Process the file (assumes PDF for simplicity)
        if file.filename.endswith('.pdf'):
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:	
                resume_text = ""
                for page in pdf.pages:
                    resume_text += page.extract_text()
                    logger.info(f'resume text is',{resume_text})

        elif file.filename.endswith('.docx'):
            # Logic for handling docx files (you can use python-docx for this)
            resume_text = docx_parser_function(file_content)
            logger.info(f'resume text is {resume_text}')# Define a custom function for parsing .docx
            
        elif file.filename.endswith('.doc'):
            print('This is a doc file')
            try:
                # try:
                #     with open(f"{file.filename}", "wb") as buffer:
                #         shutil.copyfileobj(file.file, buffer)
                # except Exception as e:
                #     return JSONResponse(content={"message": f"Error: {str(e)}"}, status_code=400)

                directory_location='/home/ragfastapi/public_html/docx_files'
                file_path=f'/home/ragfastapi/public_html/{file.filename}'
                command = [
                    'libreoffice', 
                    '--headless', 
                    '--convert-to', 'docx', 
                    file_path
                    ]
                try:
                    subprocess.run(command, check=True)
                    print('Able to convert the doc file to the docx file ..............................')
                except:
                    logger.exception('error while converting the file to the docx file')
                    traceback.print_exc()
                # subprocess.check_output(['lowriter','--convert-to','docx','--outdir',dirloc,file])
                f_name =file.filename.split('.')[0]
                print("f_name: ", f_name)
                f_name=f_name+'.docx'
                # docx_path=os.path.join(directory_location,f_name)
                resume_text = docx2txt.process(f_name)
                logger.info('Resume text is',resume_text)
                try:
                    os.remove(os.path.join(f_name))
                except:
                    traceback.print_exc()
                try:
                    os.remove(file_path.split('/')[-1])
                except:
                    traceback.print_exc()
                # return resume_text
            except:
                logger.exception()
                traceback.print_exc()
                logger.exception('error in the doc ')

        # Prepare a prompt for the Llama model to extract the necessary information
        # prompt = f"""
        # Given the following resume text, extract the full name, education, experience, and projects.

        # Resume Text:
        # {resume_text}

        # Your response should be in JSON format like:
        # {{
        # "full_name": "John Doe",
        # "education": [["Degree", "University", "Year","Grade"]],
        # "experience": [["Role", "Company", "Dates"]],
        # "projects": [["Project 1"]]
        # }}
        # do not change the output format the format and given keys should be exactely the same
        # """
        # prompt2 = f"""
        # Given the following resume text, extract the full name, education, experience, and projects.

        # Resume Text:
        # {resume_text}

        # Your response should be in JSON format like:
        # {{
        # "full_name": "John Doe",
        # "education": [{"Degree", "University", "Year","Grade","text"}],
        # "experience": [{"DESIGNATION", "COMPANY", "Dates"}],
        # "projects": [{"Project 1", "description", "estimate value"}]
        # }}
        # do not change the keys in the json response the keys should be exacty the same and if you don't find any value for the key return "" in the value 
        # for eg Degree is key and the value should be the degree found in the resume text 
        # """
        # Query the Ollama API with Llama 3.2 model

        name_prompt= f""" given the  following resume text {resume_text} extract the name and return the output in the following format {{"name":"name in the resume"}}"""
        dic={}
        try:
            name = ollama.chat(model="llama3.2:latest", messages=[{"role": "user", "content": name_prompt}])
            dic['name']=name['message']['content']
            logger.info(f'The name is {name}')
            
        except Exception as e:
            logger.exception('an erroe occured while running the model')
            raise HTTPException(status_code=500, detail="Error querying Llama model: " + str(e))
        
        # education_prompt=f"""given the  following resume text {resume_text} extract the education detail  and return the output in the following format {{"education": [["Degree1", "University1", "Year1","Grade1"],[["Degree2", "University2", "Year2","Grade2"]]]}}""" 
        education_prompt=f"""
        Given the following resume text (represented as {resume_text}), extract all the education details mentioned in the resume. For each educational entry, identify the following:

1. **Degree**: The degree or certification obtained (e.g., "Bachelor of Science," "Master of Arts," or any equivalent).
2. **University**: The name of the university or institution where the degree was earned.
3. **Year**: The year of graduation or completion of the degree.
4. **Grade**: The grade or GPA (if available), or any other indicator of academic performance (e.g., "First Class," "Magna Cum Laude," etc.).

If any detail is missing (e.g., if the grade is not provided), skip it but still include the other details if available.

Return the extracted education entries in the following format:
{{
    "education":[
        ["Degree1", "University1", "Year1", "Grade1"],
        ["Degree2", "University2", "Year2", "Grade2"]
    ]
}}
        """
        try:
            education = ollama.chat(model="llama3.2:latest", messages=[{"role": "user", "content": education_prompt}])  # Pass the prompt to the LLM
            dic['education']=education['message']['content']
            logger.info(f'The education detail extract  is {education}')
            
        except Exception as e:
            logger.exception('an erroe occured while running the model')
            raise HTTPException(status_code=500, detail="Error querying Llama model: " + str(e))
        
        # experience_prompt=f"""given the following resume text {resume_text} extract the experience and return the output in the following format {{"experience":[["role1","company1","duration1"],["role2","company2","duration2"]]}}"""
        experience_prompt=f"""
        Given the following resume text (represented as {resume_text}), extract all the professional experience entries mentioned in the resume. For each experience, identify the individual's job role, the company they worked at, and the duration of the role (e.g., months or years). Be sure to capture the most relevant and specific details, avoiding vague or unclear descriptions.

For each experience entry, the model should extract:
1. **Role**: The job title or role the individual held.
2. **Company**: The name of the company or organization where the individual worked.
3. **Duration**: The time period during which the individual worked in that role (e.g., "Jan 2020 - Dec 2021" or "3 years").

Return the extracted experience entries in the following format:
{{
    "experience":[
        ["role1", "company1", "duration1"],
        ["role2", "company2", "duration2"]
    ]
}}
        """

        try:
            experience = ollama.chat(model="llama3.2:latest", messages=[{"role": "user", "content": experience_prompt}])  # Pass the prompt to the LLM
            dic['experience']=experience['message']['content']
            logger.info(f'The experience extracted is {experience}')
            
        except Exception as e:
            logger.exception('an erroe occured while running the model')
            raise HTTPException(status_code=500, detail="Error querying Llama model: " + str(e))
        
        # project_prompt=f"""given the following resume text {resume_text} extract the project and return the output in the following format {{"project":[["project1"],["project2"],["project3"]]}}"""
        # project_prompt=f"""Given the following resume text (represented as {resume_text}), extract all the distinct projects mentioned in the resume. A "project" can be any initiative, task, assignment, or endeavor the individual has worked on, including professional work, research projects, freelance tasks, or academic contributions. The extracted projects should only contain the project names (i.e., titles or key identifiers of the projects), not job titles, descriptions, or skill listings. The projects should be returned in the following format, where each project name is enclosed in double quotation marks.Return the output in the following format:{{"project":[["project1"],["project2"],["project3"]]}}
        project_prompt=f"""
        Given the following resume text (represented as {resume_text}), extract all the distinct projects mentioned in the resume. A "project" can be any initiative, task, assignment, or endeavor the individual has worked on, including professional work, research projects, freelance tasks, or academic contributions.

For each project, extract the following details:
1. **Project Name**: The title or key identifier of the project (e.g., "AI Model Development" or "Website Redesign").
2. **Client** (if provided): The client or company associated with the project (e.g., "XYZ Corp" or "Nonprofit Org").
3. **Employer**: The company or organization the individual was employed by while working on the project (e.g., "ABC Solutions").
4. **Designation**: The role or position held by the individual during the project (e.g., "Project Manager," "Software Developer").

The extracted data should only contain these four pieces of information and avoid job titles, descriptions, or skill listings unrelated to the project. If any of the information is missing (e.g., client, employer, or designation), the entry can be left blank, but the other details should still be included.

Return the output in the following format only:

{{
    "project":[
        ["project1", "client1", "employer1", "designation1"],
        ["project2", "client2", "employer2", "designation2"],
        ["project3", "client3", "employer3", "designation3"]]
}}
        """

        try:
            project = ollama.chat(model="llama3.2:latest", messages=[{"role": "user", "content": project_prompt}])  # Pass the prompt to the LLM
            dic['project']=project['message']['content']
            logger.info(f'The project extracted is {project}')  
                        
        except Exception as e:
            logger.exception('an erroe occured while running the model')
            raise HTTPException(status_code=500, detail="Error querying Llama model: " + str(e))

        # Parse the result (assuming the response is JSON-formatted as expected)
        try:
            
            #extracted_data = result['message']['content'].replace('Here is the extracted information in JSON format:\n\n', '').replace(' // No projects mentioned in the given text\n', '')# The response should contain the structured JSON text
            # Regex patterns to extract required fields
            response = dic
            return response
            
        except ValueError:
            logger.exception()
            raise HTTPException(status_code=500, detail="Invalid response from Llama model")
    except:
        logger.exception('exception in the main function')
    # Return the extracted information as JSON
    
    
    
def docx_parser_function(file_content):
    # Load the .docx file from bytes
    doc = docx.Document(io.BytesIO(file_content))
    
    # Initialize a list to hold the extracted text
    full_text = []

    # Extract text from each paragraph in the document
    for para in doc.paragraphs:
        if para.text:
            full_text.append(para.text)

    # Extract text from tables (if any)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            full_text.append(" | ".join(row_text))

    # Join all the extracted text into a single string, separated by line breaks
    return "\n".join(full_text)
    
    
    
@app.post("/pdf-to-html", response_class=HTMLResponse)
async def upload_pdf(request: Request, file: UploadFile = File(...)):
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a PDF file.")

    # Define a local directory for storing files
    dest_dir = "./temp_files"
    os.makedirs(dest_dir, exist_ok=True)  # Ensure the directory exists

    # Save the uploaded PDF temporarily
    temp_pdf_path = os.path.join(dest_dir, file.filename)
    with open(temp_pdf_path, "wb") as temp_pdf:
        temp_pdf.write(await file.read())
    
    # Define the HTML file name
    new_file_name = f"{str(uuid4())}.html"
    html_filename = os.path.join(dest_dir, new_file_name)

    try:
        # Convert PDF to HTML using pdf2htmlEX command
        subprocess.run(
            ["pdf2htmlEX", temp_pdf_path, html_filename, "--embed-external-font", "0",  "--tounicode", "1"],
            capture_output=True,
            check=True,
        )
    except subprocess.CalledProcessError as e:
        os.remove(temp_pdf_path)
        raise HTTPException(status_code=500, detail=f"Failed to convert PDF to HTML: {e}")

    # Clean up the temporary PDF file
    os.remove(temp_pdf_path)

    download_link = str(request.url_for("download_file", file_name=new_file_name))

    # Return the download link as a JSON response
    return JSONResponse(content={"download_link": download_link})


@app.get("/download/{file_name}")
async def download_file(file_name: str, background_tasks: BackgroundTasks):
    # Correct the file path to the relative project directory
    file_path = f"./temp_files/{file_name}"

    # Check if the file exists
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    # Schedule file deletion after response to avoid clutter
    background_tasks.add_task(os.remove, file_path)

    # Return the file as an attachment
    return FileResponse(
        path=file_path,
        media_type="text/html",
        filename=file_name,
        headers={"Content-Disposition": f"attachment; filename={file_name}"}
    )


class TextRequest(BaseModel):
    text: str

@app.post("/send-text/")
async def receive_text(request: TextRequest):
    experience_prompt=f"""
    Given the following resume text (represented as {request.text}), extract all the professional experience entries mentioned in the resume. For each experience, identify the individual's job role, the company they worked at, and the duration of the role (e.g., months or years). Be sure to capture the most relevant and specific details, avoiding vague or unclear descriptions.

    For each experience entry, the model should extract:
    1. **Role**: The job title or role the individual held.
    2. **Company**: The name of the company or organization where the individual worked.
    3. **Duration**: The time period during which the individual worked in that role (e.g., "Jan 2020 - Dec 2021" or "3 years") convert the duration in the database data format YYYY-MM-DD if day is not sepcified then take day as 01.

    Return the extracted experience entries in the following format:
    {{
        "experience":[
            ["role1", "company1", "duration1"],
            ["role2", "company2", "duration2"]
        ]
    }}
            """
    logger=get_logger()
    dic={}
    try:
        name = ollama.chat(model="llama3.2:latest", messages=[{"role": "user", "content": experience_prompt}])
        dic['name']=name['message']['content']
        logger.info(f'The name is {name}')
        
    except Exception as e:
        logger.exception('an erroe occured while running the model')
        raise HTTPException(status_code=500, detail="Error querying Llama model: " + str(e))
    return {"message": f"You sent: {request.text}"}