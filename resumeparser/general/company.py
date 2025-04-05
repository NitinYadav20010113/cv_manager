import pandas as pd
from docx.api import Document
from general.experience_keywords import headings, company_keywords, designation_keywords, some_keywords, experience_keywords, project_duration_keyword, to_keyword, from_keyword, sign_keyword
import re
import spacy
from general.tblmodels import Tbldesignation, Tblcompany_master
from datetime import datetime
from nltk import ngrams
from stop_words import get_stop_words
from nltk.corpus import stopwords
from datetime import datetime, timedelta
from datefinder import find_dates
from datetime import datetime



# print("company_start---------------------------------------------------------------------------------------------")
stop_words = list(get_stop_words('en'))         #Have around 900 stopwords
nltk_words = list(stopwords.words('english'))   #Have around 150 stopwords
stop_words.extend(nltk_words)
nlp = spacy.load('en_core_web_sm')

items_to_remove = ['of', 's', 'a', 'my', 'myself']
stop_words = [i for i in stop_words if i not in items_to_remove]



def subset_rem(list_):
    # # print("list_: ", list_) 
    unique_list = []
    for item in list_:
        if 'employer' in item:
            item = item.replace('employer', '').strip()

        if not any(item is u for u in unique_list):
            unique_list.append(item)
        else:
            pass

    if 'pvt ltd' in unique_list:
        unique_list.remove('pvt ltd')
 
    return unique_list


# function for remove duplicacy of the values

def remove_des_list(item, new_item_list):
    for des in des_list:
        if len(des.split())==2:
            if des in new_item_list:
                item = item.replace(des, '').strip()
                break
    return item

def company_subset_rem(list_, text):
    # # print("list_: ", list_) 
    text = text.lower()
    month_name = ['january','jan','february','feb','march','mar','april','apr','may','june','jun','july','jul','august','aug','september','sep','october','oct','november','nov','december','dec']
    des_list.remove('engineering consultant')
    des_list.remove('independent engineer')

    # Define helper functions
    def clean_item(item):
        item = item.replace('employer', '').replace('associate developer', '').strip()
        for month in month_name:
            item = item.replace(month, '').strip()
        return item

    def join_words(words):
        return [words[i] + " " + words[i+1] for i in range(len(words)-1)] + [words[-1]]

    # Create unique_list
    unique_list = []
    for item in list_:
        item = clean_item(item)
        words = item.split()
        new_item_list = join_words(words)

        if len(unique_list) == 0:
            item = remove_des_list(item, new_item_list)
            unique_list.append(item)
        else:
            if not any(item == u for u in unique_list):
                if not any(item.split()[:2] == u.split()[:2] for u in unique_list):
                    item = remove_des_list(item, new_item_list)
                    unique_list.append(item)

    # Filter 'pvt ltd' from unique_list
    unique_list = [i.replace('-', '').strip() if i.startswith('-') else i for i in unique_list]
    if 'pvt ltd' in unique_list:
        unique_list.remove('pvt ltd')

    item_index = {text.find(i):i for i in unique_list if text.find(i)!=-1}

    updated_unique_list = [i[1] for i in sorted(item_index.items())]
    

    # # print(updated_unique_list)
    return updated_unique_list


# function for extract designation from text and match with database..............................
def designation(text):
    text = text.replace('”', '').replace('“', '')
    designation = ''
    c, e_data= [], []
    try:
        for i in range(1,3):
            s = ngrams(text.split(), i) 
            for j in s:
                c.append(' '.join(j))
                e_data.append(sorted(c, key = len, reverse = True))

        if len(e_data)!=0:
            for com in e_data:
                e_list=[]
                for i in com:
                    if any(n in i.lower().split() for n in stop_words) == False:
                        e_list.append(i)


            for el in e_list:
                el = el.lower()
                if el in des_list:
                    designation+=el
                    break
                else:
                    pass
                
            return designation 
        else:
            return ''
    except:
        pass                                                      
            
# function for extract designation from text and match with database..............................


# function for checking formate of the companies-----------------------------
def com_format_checker(text,path):
    text = text.replace('\t', ' ')
    new_des = []
    new_designation_index = []
    new_heading_index = []
    new_headings_without_table = None
    new_company_keywords = None
    new_designation_keywords = None
    new_hed = []
    

    for x in headings:
        new_heading_keywords = re.search(x.lower()+'(\n+|\t+|\s{3,}|:|-|:-| :| - | : -){1,}', text.lower())
        if new_heading_keywords == None:
            pass
        else:
            new_hed.append(new_heading_keywords)
            new_heading_index.append(new_heading_keywords.start())
    
    if len(new_heading_index) > 1:
        if new_heading_index[0]>new_heading_index[1]:
            new_headings_without_table =  new_hed[1]
        else:
            new_headings_without_table =  new_hed[0]
    elif len(new_heading_index) == 1:
        new_headings_without_table =  new_hed[0]
    else:
        new_headings_without_table = None

    
        
    if new_headings_without_table:
        found_item = new_headings_without_table.group()
        text = text.lower()
        s_index = text.find(found_item)
        new_text = text[s_index:]

    else:
        new_text = text


    new_text = new_text.replace("(employer)", "employer")
    new_text = re.sub(' +', ' ', new_text.lower())

    for x in company_keywords:
        new_company_keywordss = re.search(x.lower()+'(\n+|\t+|\s{3,}|\s\d|\s;|:|-|:-|: -|\s:-|\s: -|\s{2,}:|\s:|\s-){1,}', new_text.lower())
        if new_company_keywordss==None:
            pass
        else:
            new_company_keywords = new_company_keywordss
            break

    for x in designation_keywords:
        new_designation_keywordss = re.search(x.lower()+'(\n+|\t+|\s{3,}|:|-|:-| :| - | : -){1,}', new_text.lower())
        if new_designation_keywordss==None:
            pass
        else:
            new_des.append(new_designation_keywordss)
            new_designation_index.append(new_designation_keywordss.start())

    if len(new_designation_index) > 1:
        if new_designation_index[0]>new_designation_index[1]:
            new_designation_keywords =  new_des[1]
        else:
            new_designation_keywords =  new_des[0]
    elif len(new_designation_index) == 1:
        new_designation_keywords =  new_des[0]
    else:
        new_designation_keywords = None

    

    text = re.sub(r'\n*\t*?:\s*\n*', ':', text.lower())

    head_list, company_list, designation_list = [], [], []
    found_item = None
    company_item = None
    designation_item = None

    # code for docx file----------------------------------------------------------------------------------------

    if path.lower().endswith('.docx'):
        document = Document(path)
        table = document.tables
        text = text.lower()

        if len(table) == 0:

            if new_headings_without_table:
                found_item = new_headings_without_table.group()
                head_list.append(found_item)
            else:
                head_list = []

            if new_company_keywords:
                company_item = new_company_keywords.group()
                company_list.append(company_item)
            else:
                company_list = []
                    
            if new_designation_keywords:
                designation_item = new_designation_keywords
                designation_list.append(designation_item)
            else:
                designation_list = []

                
            if head_list:
                if company_list:
                    if designation_list:
                        data=company_without_table(text)  # function where company keywords are available
                        return data 
                    else:
                        data=ADD_PROFESSIONAL_EXPERIENCE(path, text) # functions where no any keywords are available
                        return data
                else:
                    data=ADD_PROFESSIONAL_EXPERIENCE(path, text)
                    return data
            elif company_list:
                if designation_list:
                    data=company_without_table(text)
                    return data
                else:
                    data=ADD_PROFESSIONAL_EXPERIENCE(path, text)
                    return data
            else:
                data = ADD_PROFESSIONAL_EXPERIENCE(path, text)
                return data

        elif len(table) != 0: 
           
            if new_company_keywords:
                company_list.append(new_company_keywords.group())      
            else:
                company_list = []
            
            if new_headings_without_table:
                head_list.append(new_headings_without_table.group())
            else:
                head_list = []

            if new_designation_keywords:
                designation_list.append(new_designation_keywords.group())
            else:
                designation_list = []

                
            if head_list:
                if company_list:
                    if designation_list:    
                        # print('the format is the First fomat')
                        data=company_without_table(text)
                        return data 
                    else:
                        data=ADD_PROFESSIONAL_EXPERIENCE(path, text)
                        return data
                else:
                    data=ADD_PROFESSIONAL_EXPERIENCE(path, text)
                    return data

            elif company_list:
                if designation_list:
                    data=company_without_table(text)
                    return data
                else:
                    data=ADD_PROFESSIONAL_EXPERIENCE(path, text)
                    return data

            else:
                data=ADD_PROFESSIONAL_EXPERIENCE(path, text)
                return data
    
    # code for pdf file----------------------------------------------------------------------------------------

    elif path.lower().endswith('.pdf'):

        if new_headings_without_table:
            found_item = new_headings_without_table.group()
            head_list.append(found_item)
        else:
            head_list = []
               
        if new_company_keywords:
            company_item = new_company_keywords.group()
            company_list.append(company_item)
        else:
            company_list = []

        if new_designation_keywords:
            designation_item = new_designation_keywords.group()
            designation_list.append(designation_item)
        else:
            designation_list = []

        if head_list:
            if company_list:
                if designation_list:
                    data=company_without_table(text)
                    if data[0]['COMPANY'] == '':
                        data=ADD_PROFESSIONAL_EXPERIENCE(path, text)
                        return data
                    else:
                        return data
                    
                else:
                    data=ADD_PROFESSIONAL_EXPERIENCE(path, text)
                    return data
            else:
                data=ADD_PROFESSIONAL_EXPERIENCE(path, text)
                return data
                
                
        elif company_list:
            if designation_list:
                data=company_without_table(text)
                return data
            else:
                data=ADD_PROFESSIONAL_EXPERIENCE(path, text)
                return data
        else:
            data = ADD_PROFESSIONAL_EXPERIENCE(path, text)
            return data


# end of the code........................................................................................................


#functions for with table................................................................................

def table_no_finder(path,textt):
    count=0
    if path.lower().endswith('.docx'):
        list_of_all_headings = ['employer','employing organization','Organization','name of organization','Name of Firm', 'Employing organization and your title/position. Contact information for references','Employing organization and your title/position, Contact information for references','Employing organization and your title/position, Contact  information for reference','position held','position','position held','Role','positions held','position/s held','position','Designation', 'From','To','Dates', 'From date', 'To Date', 'Duration of service', 'time period','Period', '\nPeriod\n','Employing', 'Tenure of Service', 'Employer(Position Held)', 'Employer \n\n(Position Held)']
        document = Document(path)
        table = document.tables

        for x in range(len(table)):
            table =document.tables[x]
            
            data = []

            keys = None
            for i, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)

                if i == 0:
                    keys = tuple(text)
                    continue
                row_data = dict(zip(keys, text))
                data.append(row_data)


            df = pd.DataFrame(data)
            list_of_column=list(df.columns)

            for y in list_of_column:
                for z in list_of_all_headings:
                    if z.lower() == y.lower():  
                        count+=1
                        break

            if count < 2:
                pass
            else:
                break

        if count >= 2:
            data=table_data_extractor(document,x,textt)   
            return data   
        else:
            data = com_format_checker(textt, path)
           
            for i in data:
                if i['COMPANY'] == '':
                    if i['DESIGNATION'] == '':
                        if i['FROM DATE'] == '':
                            if i['TO DATE'] == '':
                                data = ADD_PROFESSIONAL_EXPERIENCE(path, textt)
                                return data   
                    else:
                        data = ADD_PROFESSIONAL_EXPERIENCE(path, textt)
                    return data   
                return data
                
    else:
        data = com_format_checker(textt, path)
        return data


#functions for with table................................................................................

        
# function for data extracting from table
def table_data_extractor(document,table_no,textt):
    
    table = document.tables[table_no]
    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)

    # these all are keywords for getting data from table
    
    date_heading=['from','to','dates', 'period','\nperiod\n','from date', 'to date', 'tenure of service', 'duration', 'tenure of service', 'duration of service', 'date of joining', 'time period']
    from_heading=['from', 'from date']
    to_heading=['to', 'to Date']
    organisation_heading=['company employer','organization','employer (sub)','employer','company',	'employing organization',
	    'name of organization','employing organization and your title/position. contact information for references',
	    'employing organization and your title/position, contact information for references',
	    'employing organization and your title/position, contact information for reference',
        'employer (position held)','organisational history','past organization', 'name of firm', 'employing']
    
    position_held_heading=['position held','position','positions','designation','positions held',
                   'position/s held','job profile','present position','post','role']
    froms=[]
    to=[]
    organisation = []
    position_held = []

    from_dates, to_dates, match_list, main_matches_list, new_dates = [], [], [], [], []
    
    df = pd.DataFrame(data)
    df.reset_index(drop=True, inplace=True)
    list_of_headings=list(df.columns)
    
    for x in list_of_headings:
        for i in from_heading:
            match_date=re.search(i,x.lower())
            if match_date!=None:
                break
        if match_date!=None:
            froms.extend(df[x].to_list())

            # function for convert date into datetime formate
            for date_str in froms:
                a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'current']  
                if date_str.lower() in a:
                    datet = datetime.now().strftime("%Y-%m-%d")
                    from_dates.append(datet)
                else:
                    date_str = re.sub(r"[,;@#?!&$–/%^\s.’-]",' ', date_str)
                    try:
                        for i in formats:
                            try:
                                date_obj = datetime.strptime(date_str, i)
                                date = date_obj.strftime("%Y-%m-%d")
                                from_dates.append(date)
                                break
                            except:
                                pass
                    except:
                        pass

        for i in to_heading:
            match_date=re.search(i,x.lower())
            if match_date!=None:
                break
        if match_date!=None:
            to.extend(df[x].to_list())  

            # function for convert date into datetime formate
            for date_str in to:
                a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'current']  
                if date_str.lower() in a:
                    datet = datetime.now().strftime("%Y-%m-%d")
                    to_dates.append(datet)
                else:
                    date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                    try:
                        for i in formats:
                            try:
                                date_obj = datetime.strptime(date_str, i)
                                date = date_obj.strftime("%Y-%m-%d")
                                to_dates.append(date)
                                break
                            except:
                                pass
                    except:
                        pass

        if len(froms) == 0 and len(to) == 0:
            for i in date_heading:
                match_date=re.search(i,x.lower())
                if match_date!=None:
                    break
            if match_date!=None:
                froms.extend(df[x].to_list())

                # function for convert date into datetime formate
                for date_str in froms:
                    date_str = re.sub(r"[,;@#?!&$/%^\s.\n–’-]",' ', date_str).replace('to', '')
                    date = re.finditer(r'(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sep|october|oct|november|nov|december|dec|\d{2}){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sep|october|oct|november|nov|december|dec|\d{2}){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sep|october|oct|november|nov|december|dec|\d{2}|till now|till today|at present|till current date|till date|still continuing|present|continuing|current){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sep|october|oct|november|nov|december|dec|\d{2}|till now|till today|at present|till current date|till date|still continuing|present|continuing|current)?',date_str.lower())
                    for i in date:
                        matches1 = i.group(1) +' '+ i.group(2)
                        try:
                            if i.group(4):
                                matches2 = i.group(3) +' '+ i.group(4)
                            else:
                                matches2 = i.group(3)
                        except:
                            matches2 = i.group(3)

                        match_list.append(matches1)
                        match_list.append(matches2)
                        main_matches_list.append(match_list)
                        match_list = []
                
                for main_date in main_matches_list:
                    for date_str in main_date:
                        a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'current']  
                        if date_str in a:
                            datet = datetime.now().strftime("%Y-%m-%d")
                            new_dates.append(datet)
                        else:
                            date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                            try:
                                for i in formats:
                                    try:
                                        date_obj = datetime.strptime(date_str, i)
                                        date = date_obj.strftime("%Y-%m-%d")
                                        new_dates.append(date)
                                        break
                                    except:
                                        pass
                            except:
                                pass

        else:
            pass

        for i in organisation_heading:
            match_date=re.search(i.lower(),x.lower())
            if match_date!=None:
                break
        if match_date!=None:
            organisation.extend(df[x].to_list())
        for i in position_held_heading:   
            match_date=re.search(i.lower(),x.lower())
            if match_date!=None:
                break
        if match_date!=None:
            position_held.extend(df[x].to_list())

          
    f_dict={}
    main_list=[]
    
    for x in range(len(organisation)):
        try:
            f_dict['COMPANY']=organisation[x]
        except:
            f_dict['COMPANY']=''
        try:
            f_dict['DESIGNATION']=position_held[x]
        except:
            f_dict['DESIGNATION']=''
        try:    
            if len(froms)!=0 and len(froms)==len(organisation) and len(to)!=0 and len(to)==len(organisation):
                f_dict['FROM DATE']=from_dates[x]
                f_dict['TO DATE']=to_dates[x]
            else:

                # date=date_parser(textt)
                f_dict['FROM DATE'] = new_dates[x+x]
                f_dict['TO DATE'] = new_dates[x+x+1]
        except:
            f_dict['FROM DATE']=''
            f_dict['TO DATE']=''
       
        main_list.append(f_dict)
        f_dict={}
    return main_list  # main list which contain all data
    

# end of functions................................................................................

comp_list =  set(Tblcompany_master.objects.using('table_db').values_list('company_name', flat=True))
comp_list=[re.sub(r"[,;@#?!&$/%^\s.]+ *"," ",x.lower()).rstrip() for x in comp_list]


des_list =  set(Tbldesignation.objects.using('table_db').values_list('designation', flat=True))
des_list=[re.sub(r"[,;@#?!&$/%^\s.]+ *"," ",x.lower()).rstrip() for x in des_list]


def is_date(word):
    try:
        datetime.strptime(word, '%Y-%m-%d')
        return True
    except ValueError:
        return False
    
#functions for without headings................................................................................
def ADD_PROFESSIONAL_EXPERIENCE(path, text):
    new_exp, new_exp_index = [], []
    some_keyword = [s.lower() for s in some_keywords] # text ending keywords
    experience_keyword = [e.lower()+"\n" for e in experience_keywords] # text started keywords
    new_experience = [e.lower() for e in experience_keywords] 

    for x in new_experience:
        new_experience_keywords = re.search(x.lower()+'(\s?\n+|\t+|\s{3,}|:|-|:-| :| - | : -){1,}', text.lower())
        if new_experience_keywords == None:
            pass
        else:
            new_exp.append(new_experience_keywords)
            new_exp_index.append(new_experience_keywords.start())
    
    if len(new_exp_index) > 1:
        if new_exp_index[0]>new_exp_index[1]:
            new_experience_keywords =  new_exp[1]
        else:
            new_experience_keywords =  new_exp[0]
    elif len(new_exp_index) == 1:
        new_experience_keywords =  new_exp[0]
    else:
        new_experience_keywords = None

    companies_list, dates =[], []

    try:
        exp_data=""

        data=text
        data = data.replace('  ', ' ').lower()
        if any(n in data for n in experience_keyword)==True:  
            check_l, check_k=[], []
            found_item = None
            some_found_item = None

            min_index = len(text) + 1
            for i in experience_keyword:
                i= i.lower()
                text = text.lower()
                index = text.find(i)
                if index!=-1 and index<min_index:
                    found_item = i
                    min_index = index

            if found_item != None:
                check_l.append(found_item)
                s_index = text.find(found_item)
                s_data = text[s_index:]
            else:
                if new_experience_keywords:
                    found_item = new_experience_keywords.group()
                    check_l.append(found_item)
                    s_index = text.find(found_item)
                    s_data = text[s_index:]


            min_index = len(s_data) + 1
            for j in some_keyword:
                j= j.lower()
                index = s_data.find(j)
                if index!=-1 and index<min_index:
                    some_found_item = j
                    min_index = index

            if some_found_item != None:
                check_k.append(some_found_item)
            else:
                check_k = []
            
            if len(check_k) != 0:
                e_index = s_data.index(check_k[0])
                exp_data+=s_data[:e_index]
            else:
                exp_data+=s_data

        elif new_experience_keywords:
            check_l, check_k=[], []
            found_item = None
            some_found_item = None

            found_item = new_experience_keywords.group()
            check_l.append(found_item)
            s_index = text.find(found_item)
            s_data = text[s_index:]


            min_index = len(s_data) + 1
            for j in some_keyword:
                j= j.lower()
                index = s_data.find(j)
                if index!=-1 and index<min_index:
                    some_found_item = j
                    min_index = index

            if some_found_item != None:
                check_k.append(some_found_item)
            else:
                check_k = []
            
            if len(check_k) != 0:
                e_index = s_data.index(check_k[0])
                exp_data+=s_data[:e_index]
            else:
                exp_data+=s_data

        else:
            exp_data+=text[:]
        
        c=[]  
        find_dates_exp_data = re.sub(r"[,(;’.'(@#?&!$/%^:)`)]+ *"," ", exp_data)
        find_dates_exp_data = find_dates_exp_data.replace('-', 'to')
        company_exp_data = re.sub(r"[,;.(@#?&!$/%^)`–]+ *"," ", exp_data).replace('my', '').replace('crores', '')
        exp_data = re.sub(r"[,(;’.(@#?&!$/%^:)`)–]+ *"," ", exp_data).replace('c.s', 'cs')
        
        exp_data = exp_data.lower()
        find_dates_exp_data = find_dates_exp_data.lower()
        exp_data = re.sub(' +', ' ', exp_data)
        find_dates_exp_data = re.sub(' +', ' ', find_dates_exp_data)
        find_dates_exp_data = re.sub(r'\s+', ' ', find_dates_exp_data).lower()
        exp_data_new = exp_data.split('\n')            


        month_name = ['january','jan','february','feb','march','mar','april','apr','may','june','jun','july','jul','august','aug','september','sep','october','oct','november','nov','december','dec']
        a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'current']  
        
        if any(i in company_exp_data.split() for i in month_name) == True:
            c_e_p = company_exp_data.split()
            for i in month_name:
                if i in c_e_p:
                    c_e_p.remove(i)
            company_exp_data = ' '.join(c_e_p)

        if any(i in company_exp_data for i in a) == True:
            for i in a:
                if i in company_exp_data:
                    company_exp_data = company_exp_data.replace(i, '')

        company_exp_data_1 = company_exp_data.split()
        for i in company_exp_data_1:
            if i.lower().isdigit():
                company_exp_data.replace(i, '')
            else:
                pass


        # company_exp_data = re.sub(r'\d{1,}', '', company_exp_data)
        company_exp_data = re.sub(' +', ' ', company_exp_data)

        # all possible company ending keywords.......................................
        check = ['pvt ltd','pvt','ltd','limited','servicees', 'infratructure', 'jv', 'company', 'pvt limited', 'syndicate', 'freelancer', 'induatries', 'infrabuilt', 'inc']


        t_data = [] #temporary data
        for i in range(2,7):
            s = ngrams(company_exp_data.split(), i) 
            for j in s:
                c.append(' '.join(j))
                
        t_data.append(sorted(c, key = len, reverse = True)) # all combinations using ngrams
        companies=[] # companies list  

        t_list=[]
        # # print("t_data: ", t_data)
        if len(t_data)!=0:
            for com in t_data:
                for i in com:
                    if any(n in i.lower().split() for n in stop_words) == False:
                        if any(n in i.lower() for n in ['department','client',':','_',']',';','%','#','|','>', 'organization', 'experience'])==False:
                            if any(n in i.lower().split() for n in ['january','jan','february','feb','march','mar','april','apr','may','june','jun','july','jul','august','aug','september','sep','october','oct','november','nov','december','dec', 'year', 'month', 'date'])==False:
                                if any(n in i.lower() for n in a)==False:
                                    if any(n in i.lower()[0] for n in ['-'])==False:
                                        if any(n in i.lower()[-1] for n in ['-'])==False:
                                            if len(i.split()[-1]) != 1 and len(i.split()[-1])!=2:
                                                if len(i.split()[0]) != 1:
                                                    t_list.append(i)  


        # # print("before t_list: ", t_list)
        t_list = subset_rem(t_list)

        # # print("hello")
        # # print("t_list: ", t_list)
        for tl in t_list:
            tl = tl.lower()
            if tl in comp_list:
                companies.append(tl)
            elif tl.split()[-1] in check:
                companies.append(tl)
            else: 
                pass
               
        companies = company_subset_rem(companies, company_exp_data)  # all companies after removing duplicacy
                
        # # print("compamies: ", companies)
        exp_data2=[]
        designation_exp_data2 = []
        count = 0
        if len(companies)!=0:
            for i in range(len(companies)):
                l=[]
                m = []
                exp_data = re.sub(r'\s+', ' ', exp_data).lower().replace('my', '').replace('ammar', 'am')
                exp_data = re.sub(' +', ' ', exp_data)
                companies_d = companies[i].lower()
                companies_d = companies_d.replace('toubro', 'tubro').replace('constructions', 'construction')
                if i == 0:
                    if companies_d in exp_data:             
                        s_index= exp_data.index(companies_d)
                        if s_index-40 < 0:
                            l.append(exp_data[:s_index+150])
                            m.append(exp_data[:s_index+150])
                            designation_exp_data2.append(m)
                            exp_data2.append(l) 
                            month_name_present = [i for i in month_name if i in exp_data[:s_index+150]]
                            # Get the index of each item in month_name_present in exp_data
                            month_index = [exp_data[:s_index+150].index(i) for i in month_name_present]
                            if any(i for i in month_index if i<s_index):
                                count+=1
                            else:
                                pass
                        else:    
                            if len([i for i in find_dates(find_dates_exp_data[s_index-80:s_index])]) == 0:
                                if len([i for i in find_dates(find_dates_exp_data[s_index-40:s_index])]) != 0:
                                    if s_index-40 < 0:
                                        l.append(exp_data[:s_index+150])
                                        m.append(exp_data[:s_index+150])
                                        designation_exp_data2.append(m)
                                        exp_data2.append(l) 
                                    
                                    else:
                                        data = exp_data[s_index-55:s_index+370]
                                        if data == '':
                                            data = exp_data[s_index-40:s_index+370]
                                            l.append(data)
                                            m.append(data)
                                            designation_exp_data2.append(m)
                                            exp_data2.append(l) 
                                            count+=1
                                        else:
                                            l.append(data)
                                            m.append(data)
                                            designation_exp_data2.append(m)
                                            exp_data2.append(l) 
                                            count+=1

                                else:
                                    data = exp_data[s_index:s_index+370]
                                    designation_data = exp_data[s_index-40:s_index+370]

                                    m.append(designation_data)
                                    designation_exp_data2.append(m)
                                    l.append(data)
                                    exp_data2.append(l) 
                            else:
                                data = exp_data[s_index-80:s_index+370]
                                if data == '':
                                    data = exp_data[s_index-55:s_index+370]
                                    l.append(data)
                                    m.append(data)
                                    designation_exp_data2.append(m)
                                    exp_data2.append(l) 
                                    count+=1
                                else:
                                    l.append(data)
                                    exp_data2.append(l) 
                                    m.append(data)
                                    designation_exp_data2.append(m)
                                    count+=1
                    else:
                        l.append('')
                        exp_data2.append(l)
                        m.append('')
                        designation_exp_data2.append(m)

                else:
                    if companies_d in exp_data:             
                        s_index = exp_data.index(companies_d)
                        if all(i==''  for item in exp_data2 for i in item):
                            if len([i for i in find_dates(find_dates_exp_data[s_index-80:s_index])]) == 0:
                                if len([i for i in find_dates(find_dates_exp_data[s_index-40:s_index])]) != 0:
                                    if s_index-40 < 0:
                                        l.append(exp_data[:s_index+150])
                                        exp_data2.append(l) 
                                        m.append(exp_data[:s_index+150])
                                        designation_exp_data2.append(m)
                                    else:
                                        data = exp_data[s_index-55:s_index+370]
                                        if data == '':
                                            data = exp_data[s_index-40:s_index+370]
                                            l.append(data)
                                            exp_data2.append(l) 
                                            m.append(data)
                                            designation_exp_data2.append(m)
                                            count+=1
                                        else:
                                            l.append(data)
                                            exp_data2.append(l) 
                                            m.append(data)
                                            designation_exp_data2.append(m)
                                            count+=1
                                else:
                                    data = exp_data[s_index:s_index+370]
                                    designation_data = exp_data[s_index-40:s_index+370]
                                    m.append(designation_data)
                                    designation_exp_data2.append(m)
                                    l.append(data)
                                    exp_data2.append(l) 
                            else:
                                data = exp_data[s_index-80:s_index+370]
                                if data == '':
                                    data = exp_data[s_index-55:s_index+370]
                                    l.append(data)
                                    exp_data2.append(l) 
                                    m.append(data)
                                    designation_exp_data2.append(m)
                                    count+=1
                                else:
                                    l.append(data)
                                    exp_data2.append(l) 
                                    m.append(data)
                                    designation_exp_data2.append(m)
                                    count+=1
                        else:
                            if count != 0:
                                data = exp_data[s_index-80:s_index+370]
                                if data == '':
                                    data = exp_data[s_index-55:s_index+370]
                                    l.append(data)
                                    exp_data2.append(l) 
                                    m.append(data)
                                    designation_exp_data2.append(m)
                                else:
                                    l.append(data)
                                    exp_data2.append(l) 
                                    m.append(data)
                                    designation_exp_data2.append(m)
                            else:
                                data = exp_data[s_index:s_index+370]
                                designation_data = exp_data[s_index-40:s_index+370]
                                m.append(designation_data)
                                designation_exp_data2.append(m)
                                l.append(data)
                                exp_data2.append(l) 
                    else:
                        l.append('')
                        exp_data2.append(l)
                        m.append('')
                        designation_exp_data2.append(m)

        else:
            exp_data2.append([])
            designation_exp_data2.append([])


        temp=0
        try:
            if len(exp_data2[0])!=0:        
                if len(companies) !=0: 
                    comp = 0       
                    for i in range(len(companies)):
                        f_dict={}
                        f_dict['COMPANY'] = companies[i]
                        dates = date_parser(exp_data2[i][0]) 
                        if len(dates)!=0:
                            try:
                                if exp_data2[i][0].index(dates[0].split('-')[0])<exp_data2[i][0].index(companies[i]):
                                    length_dates = len(dates)
                                    if length_dates >= 2:
                                        if dates[0]>dates[1]:
                                            temp = dates[0]
                                            dates[0]=dates[1]
                                            dates[1]=temp

                                        f_dict['FROM DATE']=dates[0]
                                        f_dict['TO DATE']=dates[1]
                                    else:
                                        f_dict['FROM DATE']=dates[0]
                                        f_dict['TO DATE']=''
                                else:
                                    try:
                                        f_dict['FROM DATE']=dates[0]
                                        f_dict['TO DATE']=dates[1]
                                    except:
                                        f_dict['FROM DATE']=dates[0]
                                        f_dict['TO DATE']=''

                            except:
                                try:
                                    f_dict['FROM DATE']=dates[0]
                                    f_dict['TO DATE']=dates[1]
                                except:
                                    f_dict['FROM DATE']=dates[0]
                                    f_dict['TO DATE']=''

                        elif len(dates) == 1:
                            f_dict['FROM DATE']=dates[0]
                            f_dict['TO DATE']=''
                        else:
                            f_dict['FROM DATE']=''
                            f_dict['TO DATE']=''
                        f_dict['DESIGNATION']=designation(designation_exp_data2[i][0])
                        companies_list.append(f_dict)
                if len(companies_list) == 0:
                    f_dict={}
                    f_dict['COMPANY'] = ''
                    f_dict['FROM DATE'] = ''
                    f_dict['TO DATE'] = ''
                    f_dict['DESIGNATION'] = ''
                    companies_list.append(f_dict)
            if len(exp_data2[0])== 0:
                f_dict={}
                f_dict['COMPANY'] = ''
                f_dict['FROM DATE'] = ''
                f_dict['TO DATE'] = ''
                f_dict['DESIGNATION'] = ''
                companies_list.append(f_dict) 
            
            
        except:
            f_dict={}
            f_dict['COMPANY'] = ''
            f_dict['FROM DATE'] = ''
            f_dict['TO DATE'] = ''
            f_dict['DESIGNATION'] = ''
            companies_list.append(f_dict) 
    except:
        f_dict={}
        f_dict['COMPANY'] = ''
        f_dict['FROM DATE'] = ''
        f_dict['TO DATE'] = ''
        f_dict['DESIGNATION'] = ''
        companies_list.append(f_dict) 
    # # print("before: ", companies_list)
    new_companies_list = []
    if len(companies_list) != 1:
        for item in companies_list:
            Company = item['COMPANY']
            Designation = item['DESIGNATION']
            From_Date = item['FROM DATE']
            To_Date = item['TO DATE']

            if Company != '':
                if Designation == '':
                    if From_Date == '':
                        if To_Date == '':
                            new_companies_list.append(item)
                        else:
                            pass
                    else:
                        pass
                else:
                    if From_Date == '':
                        if To_Date == '':
                            new_companies_list.append(item)
                        else:
                            pass
                    else:
                        pass
            else:
                new_companies_list.append(item)
        
        new_c = []
        new_c_1 = []
        new_c_2 = []
        for i in companies_list:
            if i in new_companies_list:
                pass  # add new items to the copy
            else:
                new_c.append(i)

        for n in new_c:
            if 'pvt' in n['COMPANY']:
                n['COMPANY'].replace('pvt', 'private')
            
            if 'ltd' in n['COMPANY']:
                n['COMPANY'].replace('ltd', 'limited')

            new_c_1.append(n)
        
        for i in new_c_1:
            if i['COMPANY'] == 'private limited':
                new_c_1.remove(i)
        
        for d in new_c_1:
            if len(new_c_2)==0:
                new_c_2.append(d)
            else:
                count = 0
                for d2 in new_c_2:
                    if d['FROM DATE'] == d2['FROM DATE'] and d['TO DATE'] == d2['TO DATE']:
                        count +=1
                if count !=0:
                    pass
                else:
                    new_c_2.append(d)
        return new_c_2
        # print("after: ", new_c_2)


    else:
        # # print("before: ", companies_list)
        return companies_list
    


# end of the code....................................................................................................


def company_without_table(text):
    text = text.replace('date of join ing', 'date of joining').replace('date of \njoining / period', 'Date of Joining / Period')
    sign_desihnation_keyword = []
    for i in designation_keywords:
        for j in sign_keyword:
            new_item = i+j
            sign_desihnation_keyword.append(new_item)

    new_heading_keyword = ''
    new_some_keyword = ''
    new_des = []
    new_hed = []
    new_designation_index = []
    new_heading_index = []

    text = text.lower()
    main_list = []  # main list
    empty_list = []

    f_dict = {}
    
    min_index = len(text) + 1
    found_item=None

# main heading text find

    for x in headings:
        new_heading_keywords = re.search(x.lower()+'(\n+|\t+|\s{3,}|:|-|:-| :| - | : -){1,}', text.lower())
        if new_heading_keywords == None:
            pass
        else:
            new_hed.append(new_heading_keywords)
            new_heading_index.append(new_heading_keywords.start())

    if len(new_heading_index) > 1:
        if new_heading_index[0]>new_heading_index[1]:
            new_heading_keyword =  new_hed[1]
        else:
            new_heading_keyword =  new_hed[0]
    elif len(new_heading_index) == 1:
        new_heading_keyword =  new_hed[0]
    else:
        new_heading_keyword = None
    
        
    if new_heading_keyword:
        found_item = new_heading_keyword.group()
        s_index = text.find(found_item)
        new_data = text[s_index:]

    else:
        new_data = text

    new_data = new_data.replace('(employer)', 'employer')
# main heading text find

    new_company_key = ''
    new_designation_key = ''
    new_project_duration_key = ''
    new_some_company_key = []
    new_project_duration = []
    check = checker(new_data, new_company_key, new_designation_key)   # checking position of company and designation


    some_company_keywords = ['present organization', 'past organization', 'previous organization', 'present employment', 'previous employment', 'present employee', 'previous employers', 'previous employer']
    if 'past organization' and 'present organization' in new_data or 'previous organization' and 'present organization' in new_data or 'previous employment' and 'present employment'in new_data or 'present employee' and 'previous employers' in new_data or 'present employee' and 'previous employers' and 'previous employer' in new_data:
        for x in some_company_keywords:
            new_company_keywords = re.search(x.lower()+'(\n+|\t+|\s{3,}|\s;|:|- |: -|:-){1,}', new_data)
            if new_company_keywords==None:
                pass
            else:
                new_some_company_key.append(new_company_keywords)

    else:
        for x in company_keywords:
            new_company_keywords = re.search(x.lower()+'(\n+|\t+|\s{3,}|\s\d|\s;|:|-|:-|: -|\s:-|\s: -|\s{2,}:|\s:|\s-){1,}', new_data)
            if new_company_keywords==None:
                pass
            else:
                new_company_key = new_company_keywords
                break


    for x in designation_keywords:
        new_designation_keywords = re.search(x.lower()+'(\n+|\t+|\s{3,}|:|-|:-| :|  :| - | : -){1,}', new_data)
        if new_designation_keywords==None:
            pass
        else:
            new_des.append(new_designation_keywords)
            new_designation_index.append(new_designation_keywords.start())

    if len(new_designation_index) > 1:
        if new_designation_index[0]>new_designation_index[1]:
            new_designation_key =  new_des[1]
        else:
            new_designation_key =  new_des[0]
    else:
        new_designation_key =  new_des[0]

        

    new_d = 0
    for x in from_keyword:
        new_from_keyword = re.search(x.lower()+'(\n+|\t+|\s{3,}|:|-|:-| :| - | : -){1,}', new_data)
        if new_from_keyword:
            for x in to_keyword:
                new_to_keyword = re.search(x.lower()+'(\n+|\t+|\s{3,}|:|-|:-| :| - | : -){1,}', new_data)
                if new_to_keyword:
                    new_d = 1  
                else:
                    pass
        else:
            pass    

  
    some_duration_keywords = ['date of joining', 'date of joining / period']
    if 'date of joining' and 'date of joining / period' in new_data:
        for x in some_duration_keywords:
            new_project_duration_keyword = re.search(x.lower()+'(\n+|\t+|\s{3,}|\s;|:| :|- |: -|:-){1,}', new_data)
            if new_project_duration_keyword==None:
                pass
            else:
                new_project_duration.append(new_project_duration_keyword)

    elif new_d == 0:             
        for x in project_duration_keyword:
            new_data =new_data.replace('>:', ':').replace('>', ':')
            new_project_duration_keyword = re.search(x.lower()+'(\n+|\t+|\s{3,}| >|:| :|- | : -| :| - |:-){1,}', new_data)
            if new_project_duration_keyword:
                n = new_project_duration_keyword.group()
                if n == 'from' or n == 'from\n\n' or n == 'to' or n == 'to\n\n' or n == 'to\n' or n == 'from\n':
                    try:
                        if new_project_duration_keyword == None:
                            pass
                        else:
                            try:
                                new_project_duration_keyword = re.finditer(n+'(\d{2,}+|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sep|october|oct|november|nov|december|dec|\d{2,}){1,}', new_data)
                                new_project_duration_key = [x for x in new_project_duration_keyword]
                                break
                            except:
                                pass
                    except:
                        pass
                else:
                    new_project_duration_keyword = re.finditer(x.lower()+'(\n+|\t+|\s{3,}| >|:| :|- | : -| :| - |:-){1,}', new_data)
                    new_project_duration_key = [x for x in new_project_duration_keyword]
                    break

#  start of the entry code....................................................................................................
   
    if check:
        if check['status'] == 'company':
            company_list=check['company_start_index']
            for x in range(len(company_list)):
                s_index = company_list[x]
                try:
                    e_index=company_list[x+1]
                except:
                    e_index = s_index+1000

                company_text_new=new_data[s_index:e_index]                
                company_text_new = company_text_new.replace('\n\n\n\n', '\n').replace('\t', ' ')

                if len(new_some_company_key) == 2 or len(new_some_company_key) == 3:
                    for i in range(len(new_some_company_key)):
                        n = new_some_company_key[i]
                        c = n.group()
                        c = c.replace('\t', ' ')
                        c = re.sub(' + ', '', c)

                        comp = re.search(r'({})\s*(\n*)?(.+)'.format(re.escape(c)), company_text_new)
                        if comp!=None:
                            comp_ = comp.group(3)
                            if comp_ == ' ':
                                comp_=comp.group().strip()
                                remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                if any(i in comp_ for i in remove_keyword):
                                    for i in remove_keyword:
                                        comp_ = comp_.replace(i, '')
                                        break
                            else:
                                comp_ = comp.group(3)
                                remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                if any(i in comp_ for i in remove_keyword):
                                    for i in remove_keyword:
                                        comp_ = comp_.replace(i, '')
                                        break

                        elif comp == None:
                            comp = re.search(r'({})(\n*)?(.+)'.format(re.escape(c)), company_text_new)
                            if comp:
                                comp_ = comp.group(3)
                                remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                if any(i in comp_ for i in remove_keyword):
                                    for i in remove_keyword:
                                        comp_ = comp_.replace(i, '')
                                        break
                                if comp_ == ' ':
                                    comp_=comp.group().strip()
                                    remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                    if any(i in comp_ for i in remove_keyword):
                                        for i in remove_keyword:
                                            comp_ = comp_.replace(i, '')
                                            break
                                else:
                                    comp_ = comp.group(3)
                                    remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                    if any(i in comp_ for i in remove_keyword):
                                        for i in remove_keyword:
                                            comp_ = comp_.replace(i, '')
                                            break
                            else:
                                pass

                        else:
                            pass


                    if comp_:
                        d = new_designation_key.group()
                        d = re.sub(' + ', '', d)
                        d = d.replace('\t', ' ')


                        design = re.search(r'({})(\n*)?(.+)'.format(re.escape(d)), company_text_new)
                        if design!=None:
                            design=design.group(3)
                        else:
                            pass
                        
                        if new_d == 0:
                            if len(new_project_duration_key)>=1:
                                for i in range(len(new_project_duration_key)):
                                    p = new_project_duration_key[i]
                                    p = p.group()
                                    p = p.replace('\t', ' ')
                                    p = re.sub(' + ', '', p)
                                    new_dates = []

                                    dur = re.search(r'({})(\n*)?(.+)'.format(re.escape(p)), company_text_new)
                                    try:
                                        if dur!=None:
                                            dur = dur.group(3)
                                            dur_formate = new_formate_date(dur)
                                            date_d = dur_formate
                                            break
                                        else:
                                            pass
                                    except:
                                        pass


                            elif len(new_project_duration) == 2 or len(new_project_duration) == 3:
                                for i in range(len(new_project_duration)):
                                    n = new_project_duration[i]
                                    p = n.group()
                                    p = p.replace('\t', ' ')
                                    p = re.sub(' + ', '', p)

                                    dur = re.search(r'({})(\n*)?(.+)'.format(re.escape(p)), company_text_new)
                                    try:
                                        if dur!=None:
                                            dur = dur.group(3)
                                            dur_formate = new_formate_date(dur)
                                            date_d = dur_formate
                                            break
                                        else:
                                            pass
                                    except:
                                        pass

                            else:
                                new_dates = []
                                dur = None
                                date_d = date_parser(company_text_new)
                        else:
                            f = new_from_keyword.group()
                            f = re.sub(' + ', '', f)
                            f = f.replace('\t', ' ')

                            t = new_to_keyword.group()
                            t = re.sub(' + ', '', t)
                            t = t.replace('\t', ' ')

                            dur = None
                            date_d = None
                            new_dates = []


                            from_date = re.search(r'({})(\n*)?(.+)'.format(re.escape(f)), company_text_new)
                            to_date = re.search(r'({})(\n*)?(.+)'.format(re.escape(t)), company_text_new)

                            if from_date!=None:
                                from_date=from_date.group(3)
                            else:
                                pass
                            if to_date!=None:
                                to_date=to_date.group(3)
                            else:
                                pass
                            

                            comp_date=[from_date, to_date]
                            try:
                                for date_str in comp_date:
                                    a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'current']  
                                    if date_str in a:
                                        datet = datetime.now().strftime("%Y-%m-%d")
                                        new_dates.append(datet)
                                    else:
                                        try:
                                            date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                                            try:
                                                for i in formats:
                                                    try:
                                                        date_obj = datetime.strptime(date_str, i)
                                                        date = date_obj.strftime("%Y-%m-%d")
                                                        new_dates.append(date)
                                                        break
                                                    except:
                                                        pass
                                            except:
                                                pass
                                        except:
                                            pass
                            except:
                                new_dates.append('', '')

                        if comp_ != None:
                            f_dict['COMPANY'] = comp_.strip()
                        else:
                            f_dict['COMPANY'] = ''
                        if design != None:
                            f_dict['DESIGNATION'] = design.strip()
                        else:
                            f_dict['DESIGNATION'] = ''
                        
                        try:
                            if dur != None:
                                try:
                                    f_dict['FROM DATE'] = date_d[0]
                                except:
                                    pass
                                try:
                                    f_dict['TO DATE'] = date_d[1]
                                except:
                                    pass
                            else:
                                if date_d:
                                    try:
                                        f_dict['FROM DATE'] = date_d[0]
                                    except:
                                        pass
                                    try:
                                        f_dict['TO DATE'] = date_d[1]
                                    except:
                                        pass
                                elif len(new_dates)!=0:
                                    try:
                                        f_dict['FROM DATE'] = new_dates[0]
                                    except:
                                        pass
                                    try:
                                        f_dict['TO DATE'] = new_dates[1]
                                    except:
                                        pass
                                else:
                                    f_dict['FROM DATE'] = ''
                                    f_dict['TO DATE'] = ''
                        except:
                            f_dict['FROM DATE'] = ''
                            f_dict['TO DATE'] = ''

                        main_list.append(f_dict)
                        f_dict={} 

                    else:
                        pass 
                else:
                    c = new_company_key.group()
                    c = c.replace('\t', ' ')
                    c = re.sub(' + ', '', c)

                    comp = re.search(r'({})(\n*)?(.+)'.format(re.escape(c)), company_text_new)
                    if comp!=None:
                        comp=comp.group(3)
                        if any(i for i in sign_desihnation_keyword if i.lower() in comp.split()):
                            comp = ''.join([comp[:comp.index(i.lower())] for i in sign_desihnation_keyword if i.lower() in comp.split()])
                        else:
                            comp = comp
                    else:
                        pass

                
                    d = new_designation_key.group()
                    d = d.strip()
                    # d = re.sub(' + ', '', d)
                    d = d.replace('\t', ' ')


                    design = re.search(r'({})(\n*)?(.+)'.format(re.escape(d)), company_text_new)
                    if design!=None:
                        design=design.group(3)
                    else:
                        pass
                    
                    if new_d == 0:
                        if len(new_project_duration_key)>=1:
                            for i in range(len(new_project_duration_key)):
                                p = new_project_duration_key[i]
                                p = p.group()
                                p = p.replace('\t', ' ')
                                # p = re.sub(' + ', '', p)
                                p = p.strip()
                                new_dates = []

                                dur = re.search(r'({})(\n*)?(.+)'.format(re.escape(p)), company_text_new)
                                try:
                                    if dur!=None:
                                        dur = dur.group(3)
                                        dur_formate = new_formate_date(dur)
                                        date_d = dur_formate
                                        break
                                    else:
                                        pass
                                except:
                                    pass
                        else:
                            dur = None
                            new_dates = []
                            date_d = date_parser(company_text_new)
                    else:
                        f = new_from_keyword.group()
                        f = re.sub(' + ', '', f)
                        f = f.replace('\t', ' ')

                        t = new_to_keyword.group()
                        t = re.sub(' + ', '', t)
                        t = t.replace('\t', ' ')

                        dur = None
                        date_d = None
                        new_dates = []
                        

                        from_date = re.search(r'({})(\n*)?(.+)'.format(re.escape(f)), company_text_new)
                        to_date = re.search(r'({})(\n*)?(.+)'.format(re.escape(t)), company_text_new)

                        if from_date!=None:
                            from_date=from_date.group(3)
                        else:
                            pass
                        if to_date!=None:
                            to_date=to_date.group(3)
                        else:
                            pass
                        

                        comp_date=[from_date, to_date]
                        for date_str in comp_date:
                            a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'till', 'current']  
                            if date_str in a:
                                datet = datetime.now().strftime("%Y-%m-%d")
                                new_dates.append(datet)
                            else:
                                date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                                try:
                                    for i in formats:
                                        try:
                                            date_obj = datetime.strptime(date_str, i)
                                            date = date_obj.strftime("%Y-%m-%d")
                                            new_dates.append(date)
                                            break
                                        except:
                                            pass
                                except:
                                    pass


                    if comp != None:
                        f_dict['COMPANY'] = comp.strip()
                    else:
                        f_dict['COMPANY'] = ''
                    if design != None:
                        f_dict['DESIGNATION'] = design.strip()
                    else:
                        f_dict['DESIGNATION'] = ''
                    
                    if dur != None:
                        try:
                            f_dict['FROM DATE'] = date_d[0]
                        except:
                            f_dict['FROM DATE'] = ''
                        try:
                            f_dict['TO DATE'] = date_d[1]
                        except:
                            f_dict['TO DATE'] = ''
                    else:
                        try:
                            if date_d:
                                try:
                                    f_dict['FROM DATE'] = date_d[0]
                                except:
                                    pass
                                try:
                                    f_dict['TO DATE'] = date_d[1]
                                except:
                                    pass
                            elif new_dates:
                                try:
                                    f_dict['FROM DATE'] = new_dates[0]
                                except:
                                    pass
                                try:
                                    f_dict['TO DATE'] = new_dates[1]
                                except:
                                    pass
                            else:
                                f_dict['FROM DATE'] = ''
                                f_dict['TO DATE'] = ''
                        except:
                            pass

                    main_list.append(f_dict)
                    f_dict={}  
            
        elif check['status'] == 'position':
            f_dict = {}
            position_list=check['position_start_index']
            for x in range(len(position_list)):

                s_index = position_list[x]
                try:
                    e_index=position_list[x+1]
                except:
                    e_index = s_index+1000

                company_text_new=new_data[s_index:e_index] 
                date_d = date_parser(company_text_new)

                                                   
                d = new_designation_key.group()
                d = re.sub(' + ', '', d)
                d = d.replace('\t', ' ')


                design = re.search(r'({})(\n*)?(.+)'.format(re.escape(d)), company_text_new)
                if design!=None:
                    design=design.group(3)
                else:
                    pass

                if new_company_key:
                    c = new_company_key.group()
                    c = re.sub(' + ', '', c)
                    c = c.replace('\t', ' ')

                    comp = re.search(r'({})(\n*)?(.+)'.format(re.escape(c)), company_text_new)
                    if comp!=None:
                        comp=comp.group(3)
                        if any(i for i in sign_desihnation_keyword if i.lower() in comp.split()):
                            comp = ''.join([comp[:comp.index(i.lower())] for i in sign_desihnation_keyword if i.lower() in comp.split()])
                        else:
                            comp = comp

                    else:
                        pass
                else:
                    if len(new_some_company_key) == 2 or len(new_some_company_key) == 3:
                        for i in range(len(new_some_company_key)):
                            n = new_some_company_key[i]
                            c = n.group()
                            c = c.replace('\t', ' ')
                            c = re.sub(' + ', '', c)

                            comp = re.search(r'({})\s*(\n*)?(.+)'.format(re.escape(c)), company_text_new)
                            if comp!=None:
                                comp_ = comp.group(3)
                                if comp_ == ' ':
                                    comp_=comp.group().strip()
                                    remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                    if any(i in comp_ for i in remove_keyword):
                                        for i in remove_keyword:
                                            comp_ = comp_.replace(i, '')
                                            break
                                else:
                                    comp_ = comp.group(3)
                                    remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                    if any(i in comp_ for i in remove_keyword):
                                        for i in remove_keyword:
                                            comp_ = comp_.replace(i, '')
                                            break

                            elif comp == None:
                                comp = re.search(r'({})(\n*)?(.+)'.format(re.escape(c)), company_text_new)
                                if comp:
                                    comp_ = comp.group(3)
                                    if comp_ == ' ':
                                        comp_=comp.group().strip()
                                        remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                        if any(i in comp_ for i in remove_keyword):
                                            for i in remove_keyword:
                                                comp_ = comp_.replace(i, '')
                                                break
                                    else:
                                        comp_ = comp.group(3)
                                        remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                        if any(i in comp_ for i in remove_keyword):
                                            for i in remove_keyword:
                                                comp_ = comp_.replace(i, '')
                                                break
                                else:
                                    pass

                            else:
                                pass

                if new_d == 0:
                    if len(new_project_duration_key)>=1:
                        for i in range(len(new_project_duration_key)):
                            p = new_project_duration_key[i]
                            p = p.group()
                            p = p.replace('\t', ' ')
                            p = re.sub(' + ', '', p)
                            new_dates = []

                            dur = re.search(r'({})(\n*)?(.+)'.format(re.escape(p)), company_text_new)
                            try:
                                if dur!=None:
                                    dur = dur.group(3)
                                    dur_formate = new_formate_date(dur)
                                    date_d = dur_formate
                                    break
                                else:
                                    pass
                            except:
                                pass
                    else:
                        dur = None
                        new_dates = []
                        date_d = date_parser(company_text_new)
                else:
                    f = new_from_keyword.group()
                    f = re.sub(' + ', '', f)
                    f = f.replace('\t', ' ')

                    t = new_to_keyword.group()
                    t = re.sub(' + ', '', t)
                    t = t.replace('\t', ' ')

                    dur = None
                    date_d = None
                    new_dates = []


                    from_date = re.search(r'({})(\n*)?(.+)'.format(re.escape(f)), company_text_new)
                    to_date = re.search(r'({})(\n*)?(.+)'.format(re.escape(t)), company_text_new)

                    if from_date!=None:
                        from_date=from_date.group(3)
                    else:
                        pass
                    if to_date!=None:
                        to_date=to_date.group(3)
                    else:
                        pass
                    
                    
                    comp_date=[from_date, to_date]
                    for date_str in comp_date:
                        a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'current']  
                        if date_str in a:
                            datet = datetime.now().strftime("%Y-%m-%d")
                            new_dates.append(datet)
                        else:
                            date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                            try:
                                for i in formats:
                                    try:
                                        date_obj = datetime.strptime(date_str, i)
                                        date = date_obj.strftime("%Y-%m-%d")
                                        new_dates.append(date)
                                        break
                                    except:
                                        pass
                            except:
                                pass

                if comp != None:
                    f_dict['COMPANY'] = comp.strip()
                else:
                    f_dict['COMPANY'] = ''
                if design != None:
                    f_dict['DESIGNATION'] = design.strip()
                else:
                    f_dict['DESIGNATION'] = ''
                
                if dur != None:
                    try:
                        f_dict['FROM DATE'] = date_d[0]
                    except:
                         pass
                    try:
                        f_dict['TO DATE'] = date_d[1]
                    except:
                        pass
                else:
                    if date_d:
                        try:
                            f_dict['FROM DATE'] = date_d[0]
                        except:
                            pass
                        try:
                            f_dict['TO DATE'] = date_d[1]
                        except:
                            pass
                    elif new_dates:
                        try:
                            f_dict['FROM DATE'] = new_dates[0]
                        except:
                            pass
                        try:
                            f_dict['TO DATE'] = new_dates[1]
                        except:
                            pass
                    else:
                        f_dict['FROM DATE'] = ''
                        f_dict['TO DATE'] = ''


                main_list.append(f_dict)
                f_dict={}  

        else:
            f_dict = {}
            duration_list=check['duration_start_index']
            for x in range(len(duration_list)):

                s_index = duration_list[x]
                try:
                    e_index=duration_list[x+1]
                except:
                    e_index = s_index+1000

                company_text_new=new_data[s_index:e_index] 
                                                   
                d = new_designation_key.group()
                d = re.sub(' + ', '', d)

                design = re.search(r'({})(\n*)?(.+)'.format(re.escape(d)), company_text_new)
                if design!=None:
                    design=design.group(3)
                else:
                    pass
                
                if new_company_key:
                    c = new_company_key.group()
                    c = re.sub(' + ', '', c)
                    comp = re.search(r'({})(\n*)?(.+)'.format(re.escape(c)), company_text_new)
                    if comp!=None:
                        comp=comp.group(3)
                        if any(i for i in sign_desihnation_keyword if i.lower() in comp.split()):
                            comp = ''.join([comp[:comp.index(i.lower())] for i in sign_desihnation_keyword if i.lower() in comp.split()])
                        else:
                            comp = comp
                    else:
                        pass
                else:
                    if len(new_some_company_key) == 2 or len(new_some_company_key) == 3:
                        for i in range(len(new_some_company_key)):
                            n = new_some_company_key[i]
                            c = n.group()
                            c = c.replace('\t', ' ')
                            c = re.sub(' + ', '', c)

                            comp = re.search(r'({})\s*(\n*)?(.+)'.format(re.escape(c)), company_text_new)
                            if comp!=None:
                                comp_ = comp.group(3)
                                if comp_ == ' ':
                                    comp_=comp.group().strip()
                                    remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                    if any(i in comp_ for i in remove_keyword):
                                        for i in remove_keyword:
                                            comp_ = comp_.replace(i, '')
                                            break
                                else:
                                    comp_ = comp.group(3)
                                    remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                    if any(i in comp_ for i in remove_keyword):
                                        for i in remove_keyword:
                                            comp_ = comp_.replace(i, '')
                                            break

                            elif comp == None:
                                comp = re.search(r'({})(\n*)?(.+)'.format(re.escape(c)), company_text_new)
                                if comp:
                                    comp_ = comp.group(3)
                                    if comp_ == ' ':
                                        comp_=comp.group().strip()
                                        remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                        if any(i in comp_ for i in remove_keyword):
                                            for i in remove_keyword:
                                                comp_ = comp_.replace(i, '')
                                                break
                                    else:
                                        comp_ = comp.group(3)
                                        remove_keyword = ['company', 'organisation', 'organization', ':', '-', ':-', ': -']
                                        if any(i in comp_ for i in remove_keyword):
                                            for i in remove_keyword:
                                                comp_ = comp_.replace(i, '')
                                                break
                                else:
                                    pass

                            else:
                                pass

                if new_d == 0:
                    if len(new_project_duration_key)>=1:
                        for i in range(len(new_project_duration_key)):
                            p = new_project_duration_key[i]
                            p = p.group()
                            p = p.replace('\t', ' ')
                            p = re.sub(' + ', '', p)
                         
                            new_dates = []

                            dur = re.search(r'({})(\n*)?(.+)'.format(re.escape(p)), company_text_new)
                            try:
                                if dur!=None:
                                    dur = dur.group(3)
                                    dur_formate = new_formate_date(dur)
                                    date_d = dur_formate
                                    break
                                else:
                                    pass
                            except:
                                pass

                    else:
                        dur = None
                        new_dates = []
                        date_d = date_parser(company_text_new)
                else:
                    f = new_from_keyword.group()
                    f = re.sub(' + ', '', f)
                    t = new_to_keyword.group()
                    t = re.sub(' + ', '', t)
                    dur = None
                    date_d = None
                    new_dates = []


                    from_date = re.search(r'({})(\n*)?(.+)'.format(re.escape(f)), company_text_new)
                    to_date = re.search(r'({})(\n*)?(.+)'.format(re.escape(t)), company_text_new)

                    if from_date!=None:
                        from_date=from_date.group(3)
                    else:
                        pass
                    if to_date!=None:
                        to_date=to_date.group(3)
                    else:
                        pass
                    

                    comp_date=[from_date, to_date]
                    for date_str in comp_date:
                        a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'current']  
                        if date_str in a:
                            datet = datetime.now().strftime("%Y-%m-%d")
                            new_dates.append(datet)
                        else:
                            if date_str!=None:
                                date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                                try:
                                    for i in formats:
                                        try:
                                            date_obj = datetime.strptime(date_str, i)
                                            date = date_obj.strftime("%Y-%m-%d")
                                            new_dates.append(date)
                                            break
                                        except:
                                            pass
                                except:
                                    pass
                            else:
                                new_dates.append('')


                if comp != None:
                    f_dict['COMPANY'] = comp.strip()
                else:
                    f_dict['COMPANY'] = ''
                if design != None:
                    f_dict['DESIGNATION'] = design.strip()
                else:
                    f_dict['DESIGNATION'] = ''
                
                if dur != None:
                    try:
                        f_dict['FROM DATE'] = date_d[0]
                    except:
                         pass
                    try:
                        f_dict['TO DATE'] = date_d[1]
                    except:
                        pass
                else:
                    if date_d:
                        try:
                            f_dict['FROM DATE'] = date_d[0]
                        except:
                            pass
                        try:
                            f_dict['TO DATE'] = date_d[1]
                        except:
                            pass
                    elif new_dates:
                        try:
                            f_dict['FROM DATE'] = new_dates[0]
                        except:
                            pass
                        try:
                            f_dict['TO DATE'] = new_dates[1]
                        except:
                            pass
                    else:
                        f_dict['FROM DATE'] = ''
                        f_dict['TO DATE'] = ''

                main_list.append(f_dict)
                f_dict={}  
 
        return main_list  # main list which contain all data
    else:
        f_dict['COMPANY']= ''
        f_dict['DESIGNATION']=''
        f_dict['FROM DATE']= ''
        f_dict['TO DATE']=''
        empty_list.append(f_dict)
        return empty_list
    
    
#  end of the entry code....................................................................................................



# function for checking positions of company and designation........................................

def checker(new_data, new_company_keywords_value, new_designation_keywords_value):
    new_company_keyword = ''
    new_designation_keyword= ''
    new_des, new_designation_index, company_index, postition_index, duration_index, new_some_company_keyword, new_project_duration = [], [], [], [], [], [], []
    new_project_duration_keywords = ''


    some_company_keywords = ['past organization', 'present organization', 'previous organization', 'present employment','previous employment','present employee', 'previous employers', 'previous employer']
    
    if 'past organization' and 'present organization' in new_data or 'previous organization' and 'present organization' in new_data or 'previous employment' and 'present employment' in new_data or 'present employee' and 'previous employers' in new_data or 'present employee' and 'previous employers' and 'previous employer' in new_data == False:
        for x in some_company_keywords:
            new_company_keywords = re.search(x.lower()+'(\n+|\t+|\s{3,}|:|\s;|- | : -| :| - |:-){1,}', new_data)
            if new_company_keywords==None:
                pass
            else:
                new_some_company_keyword.append(new_company_keywords)

    else:
        for x in company_keywords:
            new_company_keywords = re.search(x.lower()+'(\n+|\t+|\s{3,}|\s\d|\s;|:|-|:-|: -|\s:-|\s: -|\s{2,}:|\s:|\s-){1,}', new_data)
            if new_company_keywords==None:
                pass
            else:
                new_company_keyword = new_company_keywords
                break


    for x in designation_keywords:
        new_designation_keywords = re.search(x.lower()+'(\n+|\t+|\s{3,}|:|- | : -| :|  :| - |:-){1,}', new_data)
        if new_designation_keywords==None:
            pass
        else:
            new_des.append(new_designation_keywords)
            new_designation_index.append(new_designation_keywords.start())
        
        
    some_duration_keywords = ['date of joining', 'date of joining / period']
    if 'date of joining' and 'date of joining / period' in new_data:
        for x in some_duration_keywords:
            new_project_duration_key = re.search(x.lower()+'(\n+|\t+|\s{3,}|\s;|:| :|- |: -|:-){1,}', new_data)
            if new_project_duration_key==None:
                pass
            else:
                new_project_duration.append(new_project_duration_key)
    else:
        for x in project_duration_keyword:
            new_data = new_data.replace('>:', ':').replace('>', ':')
            new_project_duration_keyword = re.search(x.lower()+'(\n+|\t+|\s{3,}| >|:| :|- | : -| :| - |:-){1,}', new_data)
            try:
                if new_project_duration_keyword:
                    n = new_project_duration_keyword.group()
                    if n == 'from' or n == 'from\n\n' or n == 'to' or n == 'to\n\n' or n == 'to\n' or n == 'from\n':
                        new_project_duration_keyword = re.search(n+'(\d{2,}+|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sep|october|oct|november|nov|december|dec|\d{2,}){1,}', new_data)
                        if new_project_duration_keyword == None:
                            pass
                        else:
                            new_project_duration_keywords = new_project_duration_keyword
                            break
                    else:
                        new_project_duration_keywords = new_project_duration_keyword
                        break
                else:
                    pass
            except:
                pass

    if new_company_keyword:
        try:
            i = new_company_keyword.group()
            co=re.finditer(i,new_data)
            for x in co:
                company_index.append(x.start())  # all company index
        except:
            pass  
    else:
        for new in new_some_company_keyword:
            try:
                i = new.group()
                i = re.sub(' + ', '', i)
                co=re.finditer(i,new_data)
                for x in co:
                    company_index.append(x.start())  # all company index
            except:
                pass


    try:
        if len(new_designation_index) > 1:
            if new_designation_index[0]>new_designation_index[1]:
                new_designation_keyword =  new_des[1]
            else:
                new_designation_keyword =  new_des[0]
        else:
            new_designation_keyword =  new_des[0]

        i=new_designation_keyword.group()
        # i = re.sub(' + ', '', i)
        i = i.strip()
        po=re.finditer(i,new_data)
        for y in po:
            postition_index.append(y.start()) # all designation index
    except:
        pass

    try:
        if new_project_duration_keywords:
            try:
                i=new_project_duration_keywords.group()
                # i = re.sub(' + ', '',  i)
                i = i.strip()
                du=re.finditer(i,new_data)
                for z in du:
                    duration_index.append(z.start()) # all duration index
            except:
                pass

            try:
                company_index.sort()
                postition_index.sort()
                duration_index.sort()
                if company_index[0]<postition_index[0]:
                    if company_index[0]<duration_index[0]:
                        data = {'status': 'company', 'company_start_index': company_index, 'position_start_index':postition_index, 'duration_start_index': duration_index}
                        return data
                    else:
                        data = {'status': 'duration', 'duration_start_index': duration_index,  'company_start_index': company_index, 'position_start_index':postition_index}
                        return data

                else:
                    if postition_index[0]<duration_index[0]:
                        data = {'status': 'position', 'position_start_index': postition_index, 'company_start_index': company_index, 'duration_start_index': duration_index}
                        return data
                    else:
                        data = {'status': 'duration', 'duration_start_index': duration_index,  'company_start_index': company_index, 'position_start_index':postition_index}
                        return data
            except:
                pass

        elif new_project_duration:
            for new in new_project_duration:
                try:
                    i = new.group()
                    i = re.sub(' + ', '', i)
                    co=re.finditer(i,new_data)
                    for x in co:
                        duration_index.append(x.start())  # all company index
                except:
                    pass
            try:
                company_index.sort()
                postition_index.sort()
                duration_index.sort()
                if company_index[0]<postition_index[0]:
                    if company_index[0]<duration_index[0]:
                        data = {'status': 'company', 'company_start_index': company_index, 'position_start_index':postition_index, 'duration_start_index': duration_index}
                        return data
                    else:
                        data = {'status': 'duration', 'duration_start_index': duration_index,  'company_start_index': company_index, 'position_start_index':postition_index}
                        return data

                else:
                    if postition_index[0]<duration_index[0]:
                        data = {'status': 'position', 'position_start_index': postition_index, 'company_start_index': company_index, 'duration_start_index': duration_index}
                        return data
                    else:
                        data = {'status': 'duration', 'duration_start_index': duration_index,  'company_start_index': company_index, 'position_start_index':postition_index}
                        return data
            except:
                pass

        else:
            try:
                company_index.sort()
                postition_index.sort()
                if company_index[0]<postition_index[0]:
                    data = {'status': 'company', 'company_start_index': company_index, 'position_start_index':postition_index}
                    return data
                else:
                    data = {'status': 'position', 'position_start_index': postition_index, 'company_start_index': company_index}
                    return data
            except:
                pass
    except:
        pass
        

#functions for checking positions of company and designation................................................................................


# new date formate start ----------------------------------------------------


def new_formate_date(text):
    text = text.replace('to ', '').replace('currunt project', 'project')
    main_matches_list = []
    new_dates = []
    match_list = []
    r = re.search(r"(\d{2}[- \.\\/]\s?\d{2}[- \.\\/]\s?\d{4})\s*?-?\s*?(till now|till today|at present|till current date|till date|still continuing|present|continuing|contd|current|\d{2}[-\.\\/]\s?\d{2}[-\.\\/]\s?\d{4})", text)    
    r_f = re.findall(r'(\d{2}[\.\/]\d{2}[\.\/]\d{2,})\s*?(\d{2}\.\d{2}\.\d{2}|till now|till today|at present|till current date|till date|still continuing|present|continuing|contd|current)', text.lower())                        
    if r:
        matches1 = r.group(1)
        matches2 = r.group(2)
        match_list.append(matches1)
        match_list.append(matches2)
        main_matches_list.append(match_list)
        match_list = []

        for main_date in main_matches_list:
            for date_str in main_date:
                a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'current']  
                if date_str in a:
                    datet = datetime.now().strftime("%Y-%m-%d")
                    new_dates.append(datet)
                else:
                    date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                    try:
                        for i in formats:
                            try:
                                date_obj = datetime.strptime(date_str, i)
                                date = date_obj.strftime("%Y-%m-%d")
                                new_dates.append(date)
                                break
                            except:
                                pass
                    except:
                        pass
        return new_dates
    
    elif len(r_f)==1:
        rf = re.finditer(r'(\d{2}[\.\/]\d{2}[\.\/]\d{2,})\s*?(\d{2}\.\d{2}\.\d{2}|till now|till today|at present|till current date|till date|still continuing|present|continuing|contd|current)', text.lower())
        for i in rf:
            matches1 = i.group(1)
            matches2 = i.group(2)
            match_list.append(matches1)
            match_list.append(matches2)
            main_matches_list.append(match_list)
            match_list = []

        for main_date in main_matches_list:
            for date_str in main_date:
                a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'current']  
                if date_str in a:
                    datet = datetime.now().strftime("%Y-%m-%d")
                    new_dates.append(datet)
                else:
                    date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                    try:
                        for i in formats:
                            try:
                                date_obj = datetime.strptime(date_str, i)
                                date = date_obj.strftime("%Y-%m-%d")
                                new_dates.append(date)
                                break
                            except:
                                pass
                    except:
                        pass
        return new_dates
        
        
        
    else:
        date_d = date_parser(text)
        return date_d
    
# new date formate end ----------------------------------------------------


#functions for main date getting code................................................................................


formats = ["%d %b %Y", "%b %Y", "%B %Y", "%b %y", "%B %y", "%m %Y", "%d %Y", "%Y %B", "%d %b", "%Y", "%B","%d%m%Y", "%d %m %Y", "%Y", "%d %B %Y", "%d %b%Y", "%B%Y", "%d %b %y", "%B %d %Y", "%b %d %Y", "%d %m %y", "%d  %m %Y"]



def date_parser(text):
    
    new_dates = []
    date_new_text = text
    date_new_text = date_new_text.replace('to', ' ').strip()
    date_new_text = re.sub(' +', ' ', text)
    text = re.sub(r'[\’,":~—-]', ' ', text).strip()
    text = text.replace('–', '').replace("'", ' ').replace('til date', 'till date').replace('present organization', '').replace('junior', '').replace('currently', '').replace('current project', 'project').replace('up to', ' ').replace(' to', ' ').replace('to[year]', ' ').replace('from[year]', ' ').replace('up', '').replace('iia', 'india').strip()
    text = re.sub(' +', ' ', text).replace("st", '').replace("nd", '').replace("rd", '').replace("th", '')
    if 'augu' in text:
        text = text.replace('augu', 'august')
    
    text = text.replace('sept', 'sep').replace('sepember', 'september')
    match_list = []
    main_matches_list = []

    d = re.findall(r'\+\s?\d{1,}[\s*]\d{1,}[\s*]?\d{1,}[\s*]?\d{1,}', text)

    for i in d:
        text = text.replace(i, '')

    e = re.findall('\S+@\S+', text)  

    for i in e:
        text = text.replace(i, '')
    
    text = re.sub(' +', ' ', text)

    date_year_values = ['ytextear', 'years', 'yrs']
    date_months_values = ['month', 'months']
    months_name = ['january','jan','february','feb','march','mar','april','apr','may','june','jun','july','jul','august','aug','september','sep','october','oct','november','nov','december','dec']
    a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'till date', 'contd', 'till', 'current']  

    split_text = text.split()
    try:
        text = [text.replace(value, '') for i, value in enumerate(split_text) if value=='till' and split_text[i+1] in months_name][0]
    except:
        text = text

    pattern = re.finditer(r'(\b\d{1,2}\s(?:january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec)\s\d{2}\b)\s*|(\b\d{1,2}\s(?:january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec)\s\d{2}\b|till now|till today|at present|till current date|till date|still continuing|present|continuing|current|till)', text.lower())
    up_pattern = re.finditer(r'((january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec)\s+(\d{1,2})\s+(\d{4}))|((january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec)\s+\d{1,2}+\s+\d{4})|(till now|till today|at present|till current date|till date|still continuing|present|continuing|current|till){1}', text.lower())
    since_pattern = re.findall(r'(since\s+\w+\s+\d{4})', text.lower())
    new_1 = re.findall(r'(\w+\s+\d{4})\s*(\d{1,2}\s+\d{2}\s+\d{4})', text.lower()) 
    new_2 = re.findall(r'(\d{1,2}\s+\d{2}\s+\d{4})\s*(\d{1,2}\s+\d{2}\s+\d{4}|till now|till today|at present|till current date|till date|still continuing|present|continuing|current|till)', text.lower())  
    len_up_pattern = len([x[0] for x in up_pattern])
    len_pattern = len([x[0] for x in pattern])

    if len(new_2) == 1:
        for main_date in new_2:
            for date_str in main_date:
                a = ['till now','till date', 'till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'till', 'current']  
                if date_str in a:
                    datet = datetime.now().strftime("%Y-%m-%d")
                    new_dates.append(datet)
                else:
                    date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                    date_str = date_str.replace('th', '').replace('sept', 'sep').replace('sepember', 'september').replace('since', '').strip()
                    try:
                        for i in formats:
                            try:
                                date_obj = datetime.strptime(date_str, i)
                                date = date_obj.strftime("%Y-%m-%d")
                                new_dates.append(date)
                                break
                            except:
                                pass
                    except:
                        pass

            return new_dates 
        
    elif len(new_1) == 1:
        for main_date in new_1:
            for date_str in main_date:
                a = ['till now','till date', 'till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'till', 'current']  
                if date_str in a:
                    datet = datetime.now().strftime("%Y-%m-%d")
                    new_dates.append(datet)
                else:
                    date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                    date_str = date_str.replace('th', '').replace('sept', 'sep').replace('sepember', 'september').replace('since', '').strip()
                    try:
                        for i in formats:
                            try:
                                date_obj = datetime.strptime(date_str, i)
                                date = date_obj.strftime("%Y-%m-%d")
                                new_dates.append(date)
                                break
                            except:
                                pass
                    except:
                        pass
            return new_dates 
        
    elif len(since_pattern) == 1:
        since_pattern.append('present')
        for date_str in since_pattern:
            a = ['till now','till date', 'till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'till', 'current']  
            if date_str in a:
                datet = datetime.now().strftime("%Y-%m-%d")
                new_dates.append(datet)
            else:
                date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                date_str = date_str.replace('th', '').replace('sept', 'sep').replace('sepember', 'september').replace('since', '').strip()
                try:
                    for i in formats:
                        try:
                            date_obj = datetime.strptime(date_str, i)
                            date = date_obj.strftime("%Y-%m-%d")
                            new_dates.append(date)
                            break
                        except:
                            pass
                except:
                    pass
        return new_dates 
    
    elif len_up_pattern==2:
        up_pattern_list = [x[0] for x in re.finditer(r'((january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec)\s+(\d{1,2})\s+(\d{4}))|((jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{1,2}+\s+\d{4})|(till now|till today|at present|till current date|till date|still continuing|present|continuing|current|till){1}', text.lower())]
        for date_str in up_pattern_list:
            a = ['till now','till date', 'till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'till', 'current']  
            if date_str in a:
                datet = datetime.now().strftime("%Y-%m-%d")
                new_dates.append(datet)
            else:
                date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                date_str = date_str.replace('th', '')
                date_str = date_str.replace('sept', 'sep')
                date_str = date_str.replace('sepember', 'september')
                try:
                    for i in formats:
                        try:
                            date_obj = datetime.strptime(date_str, i)
                            date = date_obj.strftime("%Y-%m-%d")
                            new_dates.append(date)
                            break
                        except:
                            pass
                except:
                    pass
        return new_dates
        
    elif len_pattern==2:
        pattern_list = [x[0] for x in re.finditer(r'((\d{1,2})\s+((january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec))\s+(\d{2})|(\d{1,2})\s+((jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec))\s+(\d{2}))|(till now|till today|at present|till current date|till date|still continuing|present|continuing|current|till){1}', text.lower())]
        for date_str in pattern_list:
            a = ['till now','till date', 'till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'till', 'current']  
            if date_str in a:
                datet = datetime.now().strftime("%Y-%m-%d")
                new_dates.append(datet)
            else:
                date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                date_str = date_str.replace('th', '')
                date_str = date_str.replace('sept', 'sep')
                date_str = date_str.replace('sepember', 'september')
                try:
                    for i in formats:
                        try:
                            date_obj = datetime.strptime(date_str, i)
                            date = date_obj.strftime("%Y-%m-%d")
                            new_dates.append(date)
                            break
                        except:
                            pass
                except:
                    pass
        return new_dates

    else:
        dates = re.findall(r'(\b\d{1,2}(?:st|nd|rd|th)\s+(?:january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec){1}\s?\d{4}\b)|(\btill now|till today|at present|till current date|till date|still continuing|present|continuing|current|till\b)', text.lower())
        if dates:
            if len(dates) >= 2:
                if dates[0][0] == '':
                    start_date = dates[0][1]
                else:
                    start_date = dates[0][0]
                if dates[1][0] == '':
                    end_date = dates[1][1]
                else:
                    end_date = dates[1][0]

                match_list.append(start_date)
                match_list.append(end_date)
                main_matches_list.append(match_list)
                match_list = []

                
            elif len(dates) == 1:
                if dates[0][0] not in a and dates[0][1] not in a:
                    if dates[0][0] != '':
                        start_date = dates[0][0]
                    else:
                        start_date = dates[0][1]

                    match_list.append(start_date)
                    if any(i in text.lower() for i in a) == True:
                        for i in a:
                            if i in text:
                                end_date = i
                        match_list.append(end_date)
                        main_matches_list.append(match_list)
                        match_list = []
                    else:
                        date = re.finditer(r'(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}|till now|till today|at present|till current date|till date|still continuing|present|continuing|contd|till|current){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}|till now|till today|at present|till current date|till date|still continuing|present|continuing|contd|till|current)?',text.lower())
                        for x in date:
                            if x.group(4):
                                end_date = x.group(3) +' '+ x.group(4)
                            else:
                                end_date = x.group(3)

                            match_list.append(end_date)
                            main_matches_list.append(match_list)
                            match_list = []
                    
                else:
                    date = re.finditer(r'(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}|till now|till today|at present|till current date|till date|still continuing|present|continuing|contd|till|current){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}|till now|till today|at present|till current date|till date|still continuing|present|continuing|contd|till|current)?',text.lower())
                    for x in date:

                        if x.group(1) in months_name and x.group(2) in months_name:
                            matches1 = x.group(1)
                            if x.group(3):
                                matches2 = x.group(2) +' '+ x.group(3)
                            else:
                                matches2 = x.group(2)

                            match_list.append(matches1)
                            match_list.append(matches2)
                            main_matches_list.append(match_list)
                            match_list = []
                        elif re.findall(r'\d{2}', x.group(1)) and re.findall(r'\d{2}', x.group(2)) and re.findall(r'\d{4}', x.group(3)):
                            matches1 = x.group(1) +' '+ x.group(2) +' '+ x.group(3) 
                            try:
                                if x.group(5):
                                    matches2 = x.group(4) +' '+ x.group(5)
                            except:
                                try:
                                    if x.group(4):
                                        matches2 = x.group(4)
                                    else:
                                        matches2 = ''
                                except:
                                    matches2 = ''

                            match_list.append(matches1)
                            match_list.append(matches2)
                            main_matches_list.append(match_list)
                            match_list = []

                        else:
                            matches1 = x.group(1) +' '+ x.group(2)
                            if x.group(4):
                                matches2 = x.group(3) +' '+ x.group(4)
                            else:
                                matches2 = x.group(3)

                            match_list.append(matches1)
                            match_list.append(matches2)
                            main_matches_list.append(match_list)
                            match_list = []
                
            for main_date in main_matches_list:
                for date_str in main_date:
                    a = ['till now','till date', 'till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'till', 'current']  
                    if date_str in a:
                        datet = datetime.now().strftime("%Y-%m-%d")
                        new_dates.append(datet)
                    else:
                        date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                        date_str = date_str.replace('th', '')
                        date_str = date_str.replace('sept', 'sep')
                        date_str = date_str.replace('sepember', 'september')
                        try:
                            for i in formats:
                                try:
                                    date_obj = datetime.strptime(date_str, i)
                                    date = date_obj.strftime("%Y-%m-%d")
                                    new_dates.append(date)
                                    break
                                except:
                                    pass
                        except:
                            pass
            return new_dates

        elif any(y in text for y in date_year_values) == True or any(m in text for m in date_months_values)== True and  any(mn in text for mn in months_name)== True :
            text = re.sub(r"[\n\n[year]up]",' ', text)
            date = re.finditer(r'(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}|till now|till date|till today|at present|till current date|till date|still continuing|present|continuing|contd|till|current){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}|till now|till today|till date|at present|till current date|till date|still continuing|present|continuing|contd|till|current){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}|till now|till today|till date|at present|till current date|till date|still continuing|present|continuing|contd|till|current)?',text.lower())

            for x in date:

                if x.group(1) in months_name and x.group(2) in months_name:
                    matches1 = x.group(1)
                    if x.group(3):
                        matches2 = x.group(2) +' '+ x.group(3)
                    else:
                        matches2 = x.group(2)

                    match_list.append(matches1)
                    match_list.append(matches2)
                    main_matches_list.append(match_list)
                    match_list = []
                    
                else:
                    matches1 = x.group(1) +' '+ x.group(2)
                    if x.group(5):
                        matches2 = x.group(3) +' '+ x.group(4) +' '+ x.group(5)
                    elif x.group(4):
                        matches2 = x.group(3) +' '+ x.group(4)
                    else:
                        matches2 = x.group(3)

                    match_list.append(matches1)
                    match_list.append(matches2)
                    main_matches_list.append(match_list)
                    match_list = []

            for main_date in main_matches_list:
                for date_str in main_date:
                    a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'till', 'current']  
                    if date_str in a:
                        datet = datetime.now().strftime("%Y-%m-%d")
                        new_dates.append(datet)
                    else:
                        date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                        date_str = date_str.replace('th', '')
                        try:
                            for i in formats:
                                try:
                                    date_obj = datetime.strptime(date_str, i)
                                    date = date_obj.strftime("%Y-%m-%d")
                                    new_dates.append(date)
                                    break
                                except:
                                    pass
                        except:
                            pass
            return new_dates

        elif not any(y in text for y in date_year_values) == True or not any(m in text for m in date_months_values)== True:
            text = text.replace('\n\n', '').replace('’', ' ')
                
            date = re.finditer(r'(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\b\d{2}\b)\s*(\d{4}|\bjanuary\b|\bjan\b|\bfebruary\b|\bfeb\b|\bmarch\b|\bmar\b|\bapril\b|\bapr\b|\bmay\b|\bjune\b|\bjun\b|\bjuly\b|\bjul\b|\baugust\b|\baug\b|\bseptember\b|\bsept\b|\bsep\b|\boctober\b|\boct\b|\bnovember\b|\bnov\b|\bdecember\b|\bdec\b|\b\d{2}\b)?\s*(\d{4}|\bjanuary\b|\bjan\b|\bfebruary\b|\bfeb\b|\bmarch\b|\bmar\b|\bapril\b|\bapr\b|\bmay\b|\bjune\b|\bjun\b|\bjuly\b|\bjul\b|\baugust\b|\baug\b|\bseptember\b|\bsept\b|\bsep\b|\boctober\b|\boct\b|\bnovember\b|\bnov\b|\bdecember\b|\bdec\b|\b\d{2}\b|\btill now\b|\btill today\b|\bat present\b|\btill current date\b|\btill date\b|\bstill continuing\b|\bpresent\b|\bcontinuing\b|\bcontd\b|\btill\b|\bcurrent\b)?\s*(\d{4}|\bjanuary\b|\bjan\b|\bfebruary\b|\bfeb\b|\bmarch\b|\bmar\b|\bapril\b|\bapr\b|\bmay\b|\bjune\b|\bjun\b|\bjuly\b|\bjul\b|\baugust\b|\baug\b|\bseptember\b|\bsept\b|\bsep\b|\boctober\b|\boct\b|\bnovember\b|\bnov\b|\bdecember\b|\bdec\b|\b\d{2}\b|\btill now\b|\btill today\b|\bat present\b|\btill current date\b|\btill date\b|\bstill continuing\b|\bpresent\b|\bcontinuing\b|\bcontd\b|\btill\b|\bcurrent\b)?', text.lower())    

            for x in date:
                try:
                    if re.findall(r'\d{2}', x.group(1)) and re.findall(r'\d{4}', x.group(2)):
                        try:
                            matches1 = x.group(1) +' '+ x.group(2)
                            try:
                                if x.group(4):
                                    matches2 = x.group(3) +' '+ x.group(4)
                                else:
                                    matches2 = x.group(3)
                            except:
                                matches2 = x.group(3)

                            match_list.append(matches1)
                            match_list.append(matches2)
                            main_matches_list.append(match_list)
                            match_list = []
                        except:
                            pass

                    elif re.findall(r'\d{2}', x.group(1)) and re.findall(r'\d{2}', x.group(2)):
                        pass
                    elif x.group(1) in months_name and x.group(2) in months_name:
                        matches1 = x.group(1)
                        if x.group(3):
                            matches2 = x.group(2) +' '+ x.group(3)
                        else:
                            matches2 = x.group(2)

                        match_list.append(matches1)
                        match_list.append(matches2)
                        main_matches_list.append(match_list)
                        match_list = []
                        
                    else:
                        if re.findall(r'\d{4}', x.group(1)) and re.findall(r'\d{4}', x.group(2)):
                            matches1 = x.group(1).split()[0].strip()
                            matches2 = x.group(2).  split()[0].strip()
                            match_list.append(matches1)
                            match_list.append(matches2)
                            main_matches_list.append(match_list)
                            match_list = []
                        
                        else:
                            try:
                                matches1 = x.group(1) +' '+ x.group(2)
                                try:
                                    if x.group(4):
                                        matches2 = x.group(3) +' '+ x.group(4)
                                    else:
                                        matches2 = x.group(3)
                                except:
                                    matches2 = x.group(3)

                                match_list.append(matches1)
                                match_list.append(matches2)
                                main_matches_list.append(match_list)
                                match_list = []
                            except:
                                pass
                except:
                    pass

            for main_date in main_matches_list:
                for date_str in main_date:
                    a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'till', 'current']  
                    if date_str is None:
                        pass
                    elif date_str in a:
                        datet = datetime.now().strftime("%Y-%m-%d")
                        new_dates.append(datet)
                    else:
                        date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                        date_str = date_str.replace('th', '')
                        try:
                            for i in formats:
                                try:
                                    date_obj = datetime.strptime(date_str, i)
                                    date = date_obj.strftime("%Y-%m-%d")
                                    new_dates.append(date)
                                    break
                                except:
                                    pass
                        except:
                            pass
            return new_dates

        else:
            matches = re.findall(r'(\d+)\s*(?:year(?:s)?|yr(?:s)?|yrs?)?(?:\s+and)?\s*(\d+)\s*(?:month(?:s)?|mth(?:s)?|months?)?', text.lower())
            if matches:
                for match in matches:
                    total_years = int(match[0])
                    total_months = int(match[1])
                    from_date = datetime.today() - timedelta(days=total_years*365 + total_months* 30)
                    to_date = datetime.today()

                    from_date_formatted = from_date.strftime("%Y-%m-%d")
                    to_date_formatted = to_date.strftime("%Y-%m-%d")

                    new_dates.append(from_date_formatted)
                    new_dates.append(to_date_formatted)
                    
                return new_dates
            else:
                date = re.finditer(r'(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}|till now|till today|at present|till current date|till date|still continuing|present|continuing|contd|till|current){1}\s*(\d{4}|january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|august|aug|september|sept|sep|october|oct|november|nov|december|dec|\d{2}|till now|till today|at present|till current date|till date|still continuing|present|continuing|contd|till|current)?',text.lower())
                for x in date:
                    if x.group(1) and x.group(2) in months_name:
                        matches1 = x.group(1)
                        if x.group(3):
                            matches2 = x.group(2) +' '+ x.group(3)
                        else:
                            matches2 = x.group(2)
                        match_list.append(matches1)
                        match_list.append(matches2)
                        main_matches_list.append(match_list)
                        match_list = []
                    else:
                        matches1 = x.group(1) +' '+ x.group(2)
                
                        if x.group(4):
                            matches2 = x.group(3) +' '+ x.group(4)
                        else:
                            matches2 = x.group(3)

                        match_list.append(matches1)
                        match_list.append(matches2)
                        main_matches_list.append(match_list)
                        match_list = []

                for main_date in main_matches_list:
                    for date_str in main_date:
                        a = ['till now','till today','at present','till current date', 'till date', 'still continuing', 'present', 'continuing', 'contd', 'till', 'current']  
                        if date_str in a:
                            datet = datetime.now().strftime("%Y-%m-%d")
                            new_dates.append(datet)
                        else:
                            date_str = re.sub(r"[,;@#?!&$/%^\s.’-]",' ', date_str)
                            date_str = date_str.replace('th', '')
                            date_str = date_str.replace('sept', 'sep')
                            try:
                                for i in formats:
                                    try:
                                        date_obj = datetime.strptime(date_str, i)
                                        date = date_obj.strftime("%Y-%m-%d")
                                        new_dates.append(date)
                                        break
                                    except:
                                        pass
                            except:
                                pass
                return new_dates

#functions for main date getting code................................................................................
