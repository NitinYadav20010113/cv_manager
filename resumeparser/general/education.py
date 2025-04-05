import os
import re
import docx2txt
import re, string
import nltk
from stop_words import get_stop_words
from nltk.corpus import stopwords
import spacy
from textblob import TextBlob
from nltk import ngrams
import pgeocode
from guess_indian_gender import IndianGenderPredictor
# from googlesearch import search
from difflib import SequenceMatcher
from .tblmodels import Tblinstitutes
from functools import reduce
from docx.api import Document
from selenium.webdriver.firefox.options import Options
from selenium.webdriver.firefox.service import Service as FirefoxService
from selenium import webdriver
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import traceback
from selenium.webdriver.common.by import By
from textblob import TextBlob
from .tblmodels import Tblcompany_master,Tbldesignation,Tbluniversities
import pandas as pd
from .name import name_recog
from datetime import date
import pdfplumber

stop_words = list(get_stop_words('en'))         #Have around 900 stopwords
nltk_words = list(stopwords.words('english'))   #Have around 150 stopwords
stop_words.extend(nltk_words)
nlp = spacy.load('en_core_web_sm')
stop_words = [x for x in stop_words if x!='of']
from pathlib import Path
from os import path
BASE_DIR = Path(__file__).resolve().parent.parent
flurl = str(BASE_DIR) + "/media/"



# text_data2=[]
# for file in os.listdir('C:\\Users\\raj\\Desktop\\R315\\'):
#     txt = docx2txt.process('C:\\Users\\raj\\Desktop\\R315\\'+file)
#     txt = txt.replace("\n\n","  ")
#     txt = txt.replace("\n", "  ")
#     txt = txt.replace("\t","  ")
#     text_data2.append(txt)


def pprocessor(text):
    txt = text.replace('\n\n',"  ")
    txt = txt.replace('\n', "  ")
    txt = text.replace('\t',"  ")
    return txt


# Education Degrees
PHD = ['PHD','P.HD.','P.H.D.', 'PHD.','PHD.']
POST_GRADUATION = ['M.A','M.COM', 'M.COM.','MASTERS','M TECH','M. TECH','ME','MASTER','M.E','M.E.','ME.','M.S','M.S.','M.TECH','M.TECH.','M TECH','MBA','MASTER OF ENGINEERING','M.SC.','MTECH','MA','M A','MCOM','M COM','MS','POST GRADUATION','MSC']
POST_GRADUATION = sorted(POST_GRADUATION, key = len, reverse=True)
GRADUATION = ['B.A (HONS.)','BE','BACHELORS','B E','B.E(MECH)','B.E(CIVIL)','B.E.', 'B.E','B.E ','GRADUATE','BACHELOR OF ENGINEERING','BACHELOR OF TECHNOLOGY','B.SC', 'B.SC.','C.A.','C.A','B.COM','B.COM.','B.TECH.','B.TECH', 'B TECH','B. TECH','B.A','BA','B.A.','B.S.','B.S','B. E.','B. TECH.','B. TECH','BACHELOR DEGREE','BACHELOR',"BECHALOR'S","MASTER'S","BACHALOR'S","BACHELOR'S",'B S','BS','BTECH','GRADUATION','BSC','BCOM']
GRADUATION = sorted(GRADUATION, key = len, reverse=True)
XII = ['HSC','12TH','XII','XIITH','INTERMEDIATE','GRADE–12','HSCC','HIGHER SECONDARY EDUCATION','CLASS 12','CLASS12','12 TH','12TH','HIGHER SECONDARY']
X = ['SSC','SSLC','S.S.C','10TH','MATRIC','XTH','X', 'MATRICULATION','GRADE-10','HIGH SCHOOL','10TH','10 TH','CLASS 10','CLASS10']
DIPLOMA = ['DIPLOMA','POLYTECHNIC']
CERTIFICATE_COURSE = ['CERTIFICATE']
CERTIFICATION_COURSE = ['CERTIFICATION']


university=['college','institute','school','academy','institution','university','science & technology','science','technology']
# qualification_list=['qualification','Name of Course']
# university_list=['university','University/Board','Name of board / university']
# college_list=['college /school name']
# year_of_passing_list=['year of passing','Year of completion']
# percentage_list=['percentage','CGPA/PERCENTAGE','']
# list_of_all_headings = ['Qualification','University','College /School Name','Year of Passing','Percentage','Employing organization and your title/position. Contact information for references','Employing organization and your title/position, Contact information for references','Employing organization and your title/position, Contact  information for reference','position held','position','position held','positions held','position/s held','position','From','To','Dates','Specialization, if any','Name of board / university','Name of Course','Year of completion','Name of Institute']



qualifications_list=['qualifications','examinations','class','academicqulification','qualification','exam','nameofinstitution','course','degree/diploma','degree(s)ordiploma(s)obtained','degreeobtained','standard','degreesobtained','course','degree/certificates/diploma','exam.passed','exam/degreepassed','examination','qualification','exam.','degree/examination','degree(s)/diploma(s)','level','degree','course','degree(branch)','levelofqualification','exam.passed','degree/certificate','qualification/degree','examinationpassed','particular','nameofexamination','nameofcourse','qualifications','nameofdegree','academiclevel','particulars','program','educationalqualification','nameofthedegree','detailsofeducation','degree/course','academics']
year_of_passing_list=['year','yearofpassing','yearofpassout','yearofpass','exam/year','yearofpassing','yearof','year','durationofcourse','passingyear','yearsattended','yearofpassing','datesattended','yearofpassing','yearofpassing','passingyear','yop','period','month&year','yearofdegreeobtained','yearofcompletion','month-year','yearofstudy','academicyear']
university_list=['college/school','borad/universitay','university','university/board','borad /universitay','college','board/university','college / university','board/university','university','university/board','university / board.','board/ university','institution/university','institute/university','board','examiningbody','board/uni.','institute/univ./board','boardofstudy','university/institute','awardednameofuniversity','board/council','nameofboard/university','board&uni.','universityboard','universityorboard','college&board','nameofboard/university','nameofthecourse/university','nameofinstitution&board|university','university/institutename','board/university(yearofpassing)']
collage_list=['college/institute','institute','school\college','school\college','instituteofeducation','institution','nameofcollege/university','nameofinstitution','educationalinstitutions','nameoftheinstitution','school/college','institute','nameofinstitute','school/college/university','nameofeducational institute/college&place','nameofcollege/institute','nameofcollege','institute/school','institution&place','institutionattended','frominstitution','nameofeducationalinstitution/university','nameoftheschool/college/university','institutionname&address','college/school','college/university','institute']
grade_list=['percentahe/cgpa','%ofmarks','percentage','overall%ofmarks','score','grade','percentage','percentagemarks','classobtained','percentageofmarks','%ofmarks','cgpa/%','%','marks','marks(%)','%score','classawarded','percentage/grade/cgpa','percentage/cgpa','cgpa/percentage','percentage obtained','percentage(%)','%ofmarksorcgpa','%ofpassing','remarks(%ofmarks,divisionetc)','remarks','percentage/grade','aggregate%','%marks/cpi/cgpa','result','scores/grades','%/cgpa','cgpa/%age','%obtained','percentage/division','cpi/%','marksandgrade','percentageobtained(%)','cgpa/percent','division','percentage/cgpa/cpi','grade/%marks','percentage/gradepoints','%ordivi.','aggregate','%age/cgpa','c.g.p.a/percentage','%Marks/Points']
main_list=qualifications_list+year_of_passing_list+university_list+collage_list+grade_list
# # print('the main list is',main_list)

#functions which we are using in education.......................................
   #functions which we are using in education.......................................
class Education1:
    
    # This function checks if the cv is pdf or docx and works according to that 
    def education_format_checker(institute_names,university_names,path,text,logger):
        path_check=path.split('.')[-1]
        if path_check=='pdf':
            return Education1.education_pdf_format_checker(institute_names,university_names,path,text,logger)
        elif path_check=='docx':
            return Education1.education_docx_format_cheker(institute_names,university_names,path,text,logger)
        elif path_check=='doc':
            data=Education1.qual(institute_names,university_names,text,logger)
            return data
        
    # checks if the education is present in tabular format or not for docx 
    def education_docx_format_cheker(institute_names,university_names,path,texttt,logger):
        try:
            document = Document(path)
            table = document.tables    
            for x in range(len(table)):
                try:
                    table =document.tables[x]
                    
                    data = []

                    keys = None
                    for i, row in enumerate(table.rows):
                        text = (cell.text for cell in row.cells)

                        if i == 0:
                            keys = tuple(text)
                            continue
                        row_data = dict(zip(keys, text))
                        data.append(row_data)
                    df = pd.DataFrame(data)
                    list_of_column=list(df.columns)

                    count=0
                    for y in list_of_column:
                        for z in main_list:
                            match=re.search(z.lower(),y.lower())
                            if match!=None:
                                break
                        if match!=None:
                            count+=1
                        if count>2:
                            data= Education1.education_docx_table_extractor(df)
                            return data
                        
                except:
                    logger.exception('erro in the education docx format checker function')
                    data= Education1.qual(institute_names,university_names,texttt,logger)
                    return data

        
            
            data= Education1.qual(institute_names,university_names,texttt,logger)
            return data
        except:
            logger.exception('erro in the education docx format checker function')
            data=Education1.qual(institute_names,university_names,texttt,logger)
            return data
            
        
    # Extracting the table from the doc format if the the education details are present in the tabular fomat 
    def education_docx_table_extractor(df):
        dic={}
        qualification=[]
        collage=[]
        university=[]
        year_of_passing=[]
        percentage=[]
        data = []
        main=[]

        df = df
        list_of_column=list(df.columns)
        # print('the list of columns is ',list_of_column)
        for y in list_of_column:
            for x in qualifications_list:
                a=re.search(x.lower(),y.lower())
                if a!=None:
                    qualification.extend(df[y].to_list())
            if y.lower() in collage_list:
                collage.extend(df[y].to_list())
            if y.lower() in year_of_passing_list:
                year_of_passing.extend(df[y].to_list())
            if y.lower() in grade_list:
                percentage.extend(df[y].to_list())
            if y.lower() in university_list:
                university.extend(df[y].to_list())
                
                
        for x in range(len(qualification)):
            try:    
                dic['QUALIFICATION_LEVEL']=qualification[x]
            except:
                dic['QUALIFICATION_LEVEL']=''
            try:
                dic['YEAR OF PASSING']=year_of_passing[x]
            except:
                dic['YEAR OF PASSING']=''
            try:
                dic['INSTITUTE_NAME']=collage[x]
            except:
                dic['INSTITUTE_NAME']=''
            try:
                dic['BOARD']=university[x]
            except:
                dic['BOARD']=''
            try:
                dic['GRADE']=percentage[x]
            except:
                dic['GRADE']=''
            try:
                dic['TEXT']=qualification[x]+year_of_passing[x]+collage[x]+university[x]+percentage[x] 
            except:
                dic['TEXT']=''
                
            main.append(dic)
            dic={}
        return main
            
        
            
    

    # Function for extracting the education table if the file is pdf format and the education details is in tabular format
    def education_pdf_table_extraction(df):
        # print('the data frame is',df.columns)
        qualification_details=[]
        year_of_passing_details=[]
        collage_details=[]
        universtiy_details=[]
        grade_details=[]
    #     # print('the data frame is',df)
        colums=list(df.columns)
        # print('the columns in the table are',colums)
    #     # print('the columns in the table are',colums)
        for x in colums:
            if x!=None:
                # print('the value of x is',x)
                col=re.sub(r'[\n\s\t]','',x)
                col=col.lower()
                # print('the value of the column is',col)
                if col in qualifications_list:
                    qualification_details=df[x].to_list()
                    qualification_details=[y.replace('\n','') for y in qualification_details if y!=None]
                if col in year_of_passing_list:
                    year_of_passing_details=df[x].to_list()
                if col in collage_list:
                    collage_details=df[x].to_list()
                if col in university_list:
                    university_details=df[x].to_list()
                if col in grade_list:
                    grade_details=df[x].to_list()

        # print('the year of passing details is',year_of_passing_details)        
        main=[]
        dic={}
        for x in range(len(qualification_details)):
            for y in PHD:
                sear=re.search(y,qualification_details[x].upper().replace('.',''))
                if sear!=None:
                    dic['LEVEL']='Doctorate'
            for y in GRADUATION:
                sear=re.search(y,qualification_details[x].upper().replace('.',''))
                if sear!=None:
                    dic['LEVEL']='Graduation'
            for y  in POST_GRADUATION:
                sear=re.search(y,qualification_details[x].upper().replace('.',''))
                if sear!=None:
                    dic['LEVEL']='Post Graduate'
                    
            for y  in DIPLOMA:
                sear=re.search(y,qualification_details[x].upper().replace('.',''))
                if sear!=None:
                    dic['LEVEL']='Diploma'
                    
            for y  in X:
                sear=re.search(y,qualification_details[x].upper().replace('.',''))
                if sear!=None:
                    dic['LEVEL']='10th'
            for y  in XII:
                sear=re.search(y,qualification_details[x].upper().replace('.',''))
                if sear!=None:
                    dic['LEVEL']='12th'
                    
            try:    
                dic['QUALIFICATION_LEVEL']=qualification_details[x]
            except:
                dic['QUALIFICATION_LEVEL']=''
            try:
                dic['YEAR OF PASSING']=year_of_passing_details[x]
            except:
                dic['YEAR OF PASSING']=''
            try:
                dic['INSTITUTE_NAME']=collage_details[x]
            except:
                dic['INSTITUTE_NAME']=''
            try:
                dic['BOARD']=university_details[x]
            except:
                dic['BOARD']=''
            try:
                dic['GRADE']=grade_details[x]
            except:
                dic['GRADE']=''
            try:
                dic['TEXT']=qualification_details[x]+year_of_passing_details[x]+collage_details[x]+university_details[x]+grade_details[x] 
            except:
                dic['TEXT']=''

            main.append(dic)
            dic={}

        return main

    # This function checks whether the education is inside the table or not 
    def education_pdf_format_checker(institute_names,university_names,path,text,logger):
        filename=path
        pdf=pdfplumber.open(filename)
        # # print(len(pdf.pages))

        for x in range(len(pdf.pages)):
            count=0
            table=pdf.pages[x].extract_table()
            # # print(table)
            df=pd.DataFrame(table)
            # print(len(df))
            if len(df)>1:
                try:
                    df.columns=df.iloc[0]
                    df.drop(0,inplace=True)
                    table_headings=list(df.columns)
                    for x in table_headings:
#............................... made modification to this code .....................................................................
                        if x==None:
                            continue
                        x=re.sub(r'[\n\t\s]','',x)
                        for y in main_list:
                            try:
                                if x.lower()==y.lower():
                                    count+=1
                                    break
                            except:
                                pass
                                
                        # print('the value of the count is',count)   
                        if count>=3:
                           # print('got the educational details from the table format')
                           return Education1.education_pdf_table_extraction(df)
    
                except:
                    logger.exception('erro in the education pdf format checker function')

                    
        # print('got the education details from qual function')
        return Education1.qual(institute_names,university_names,text,logger)
#     # print('the data frame is',df)

            
            
# # print('the education detais is',education_from_table_pdf())


        
    
            

    

    # this function checks which comes first the institute or the degree to avoid the mismatch to degrees and the instiutes
    def checker(qualifications,institution_indexes,year_indexes):
        """ This function checks wether the degree comes first in the education or the insitue name comes first 

        Args:
            qualifications (list): A list which contains all the degree which we foun in the resume 
            institution_indexes (list): Contains all the institues which we found in the resume 

        Returns:
            text: Its either degree or instiute depends on which come first
        """        
        # return 'year'
        try:
            year_indexes.sort()
            qualifications.sort(key=lambda x: x[2])
            institution_indexes.sort(key=lambda x:x[1])
            if institution_indexes[0][1]<qualifications[0][2] and institution_indexes[0][1]<year_indexes[0]:
                return 'institute'
            elif qualifications[0][2]<institution_indexes[0][1] and qualifications[0][2]<year_indexes[0]:
                return 'degree'
            elif year_indexes[0]<institution_indexes[0][1]  and year_indexes[0]<qualifications[0][2]:
                return 'year'
            else:
                return 'degree'
        except:
            return 'degree'
        
    # this function is to find the institute and the passing year 
    def year_and_institute(text,qual,degree,institute_names,i,s,e,tex):
        """This function finds the calls the function for finding istitue and also matches the correct passing year of the degree

        Args:
            text (str): The whole text of the resume
            qual (list): A list of qualification
            degree (list): A list of degree
            institute_names (list): _description_
            i (_type_): _description_
            s (_type_): _description_
            e (_type_): _description_
            tex (str): The text in which we need to find the institute 

        Returns:
            d(dic): the dictionary
        """        
        d={}       
        current_year=date.today()
        
        p_yr=re.finditer(r"(\d{4}(-|to|TO|To)\d{4})|(?<!\d)\d{4}(?!\d)",tex)
        p_yr=[x.group() for x in p_yr]
        for x in range(len(p_yr)):
            if len(p_yr[x])>4:
                p_yr[x]=p_yr[x][-4::]
        p_yr=[x for x in p_yr if int(x)>1900 and int(x)<current_year.year]
        institute_name=Education1.institute_name_finder(tex,institute_names,degree)
        try:
            institute_name.sort(key=lambda x:len(x))
            institute_name=institute_name[-1]
        except:
            institute_name=''
            traceback.print_exc()
        # print(institute_name)
        if len(qual)!=0:
            d['LEVEL']=qual[i]
        else:
            d['LEVEL']=''
        if len(degree)!=0:
            d['QUALIFICATION_LEVEL']=degree[i]
        else:
            d['QUALIFICATION_LEVEL']=''
        if len(p_yr)!=0:
            d['YEAR OF PASSING']=p_yr[0]
        else:
            d['YEAR OF PASSING']=''   
                    
        if institute_name=='':
            d['INSTITUTE_NAME']=''
        else:
            # # print(inst_name(text[s:e]))
            d['INSTITUTE_NAME']=institute_name
        d['BOARD']=Education1.univer_name(text[s:e])
        try:
            d['GRADE']=re.findall(r'([0-9]{1,}[.]{0,1}[0-9]{0,})[ ]{0,}%',text[s:e])[0]
        except:
            d['GRADE']=''
        d['TEXT']=text[s:e]
        
        return d
    
    # this function finds all the years in the education 
    def year_finder(text,match_index):
        """This function finds all the year which we get in the education

        Args:
            text (text): The text which we got from the resume
            match_index (int): The index of the main heading of the education...

        Returns:
            years(list): A list which contains all the years that we found in the education section of resume 
        """        
        today=date.today()
        tex=text[match_index:match_index+500]
        
        years=re.finditer(r"(\d{4}\s*(-|TO|to|To){1}\s*\d{4})|(?<!\d)\d{4}(?!\d)",tex)
        years=[x.group() for x in years]
        for x in range(len(years)):
            if len(years[x])>4:
                years[x]=years[x][-4::]
        # years=re.finditer(r'(\s|\n|\t){1}\d{4}(\s|\n|\t){1}')
        selected_years=[x for x in years if int(x)>1900 and int(x)<today.year]  
        selected_year_index=[tex.index(x) for x in selected_years]
        return selected_years,selected_year_index
                
        
  
    # The qualification_finder function finds degree in teh resume ...........
    def qualification_finder(text,index,institute_name):
        """ this function finds the degrees from the resume 

        Args:
            text (str): The whole text of the resume 
            index (int): The index of the main heading of the education from where we search for our information..
            institute_name (set): It is a set which contains all the institute from the database used set instead of list becaue time complexity of set for searching is quite good as compare to a list

        Returns:
            result(dic): A dictionary which contains all the degree which we found in that resume 

  
        """        
        institute_indexes=[]
        institute_name=[re.sub(r"[,;@#?!&$/%^\s.]+ *"," ", x.lower())  for x in institute_name]
        if index=='00':
            t_text=text
        
        else:    
            text1 = text[index:index+700]      
            t_text = text1
            
        sentence = re.sub(r"[,;@#?!&$/%^\"°-]+ *"," ", t_text)
        sentence=sentence.replace('.','')
        sentence = sentence.upper()
        grams = []
        for i in range(1,4):
            s = ngrams(sentence.split(), i) 
            for j in s:
                grams.append(' '.join(j))
        qualifications=[]
        for word in grams:
            if word in PHD:
                qualifications.append(['Doctorate',word,sentence.upper().index(word)])          
            if word in POST_GRADUATION:
                try:
                    qualifications.append(['Post Graduate',word,sentence.upper().index(word)])
                except:
                    pass                
            if word in GRADUATION:
                try:
                    qualifications.append(['Graduation',word,sentence.upper().index(word)])
                except:
                    pass 
            if word in DIPLOMA:
                try:
                    qualifications.append(['Diploma',word,sentence.upper().index(word)])
                except:
                    pass
                
            if word in XII:
                try:
                    qualifications.append(['12th',word,sentence.upper().index(word)])  
                except:
                    pass
                     
            if word in X:
                try:
                    qualifications.append(['10th',word,sentence.upper().index(word)])
                except:
                    pass
                    
            if word.lower() in institute_name or word.lower() in university:
                try:
                    institute_indexes.append([word,sentence.index(word)])
                except:
                    pass
            
        result={'qualifactions':qualifications,'sentence':sentence,'institute_indexes':institute_indexes} 
        return result

    # this is the main function this function contains calls for the other functions and this function returns
    # the final dictionary of the qualifications
    
    def qual(institute_names,university_names,text,logger):
        try:
            """ This is the main function for education

            Args:
                text (str): complete text of the cv or resume

            Returns:
                qualifin(list): a list of dictionary containing all the infomation that we need about the education
            """        
            
           
            institute_names =institute_names.union(university_names)
            qualifn=[]
            qualifications=[]
            text=' '.join([x for x in text.split() if x not in stop_words])
            list_of_education_heading=['Educational Details','Academic Details','Education','Qualification','Academic','EDUCATION DETAILS','EDUCATION','EDUCATIONAL QUALIFICATION','Educational Qualification','EDUCATIONAL CREDENTIALS','Educational Background','Academic Carrier','ACADEMIA','PERSONAL DOSSIER','Professional Qualification','ACADEMIC QUALIFICATIONS','Academic Qualifications','Professional/Academic Qualification','Summary of Qualification','Academic & Professional Qualification','EDUCATION AND QUALIFICATIONS','Academic Profile','ACADEMICS','EDUCATION PROFILE','Education Summary','Academia','ACADEMIC & TECHNICAL QUALIFICATION','Education History, Qualifications','Technical Qualification','ACADEMIC BACKGROUND','Educational Qualifications','Educational Profile','Academic Credential','SCHOLASTICS','EDUCATION BACKGROUND','qualification details','key qualifications','Education Qualification']
            match=[]
            for b in list_of_education_heading:
                match_new = re.finditer(b.lower(),text.lower())
                for x in match_new:
                    match.append(x.start())
            match=list(set(match))
            match.sort()
            all_qualification=[[]]
            for x in range(len(match)):
                match_index_check=match[x]
                data=Education1.qualification_finder(text,match_index_check,institute_names)
                qualifications=data['qualifactions']
                # print('qualifications=====2===',qualifications)
                
                
                if len(qualifications)>len(all_qualification[0]):
                    match_index=match[x]
                    all_qualification.insert(0,qualifications)
                    sentence=data['sentence']
                    institution_indexes=data['institute_indexes']
                    years=Education1.year_finder(text,match_index)
                    year_indexes=years[1]
                    years=years[0]
                    # print(years)
                    try:
                        check=Education1.checker(qualifications,institution_indexes,year_indexes)
                    except:
                        check='degree'
                    
                    
                
            qualifications=all_qualification[0]
                            
            if len(match)==0 or len(qualifications)==0:
                match='00'
                match_index=0
                data=Education1.qualification_finder(text,match,institute_names)
                qualifications=data['qualifactions']
                sentence=data['sentence']
                year_indexes=[]
                years=[]
                check='degree'
            
            try:
                for x in range(len(qualifications)):
                    for y in range(x+1,len(qualifications)):
                        if qualifications[x][1]==qualifications[y][1]:
                            qualifications.pop(y)
            except:
                logger.exception('error in the education qual function')

            
            qualifications.sort(key = lambda x: x[2])  
            # print('the qualifaction is',qualifications)   
            degree = [x[1] for x in qualifications]
            qual = [x[0] for x in qualifications]
            c=1
            
            for i in range(len(degree)):
                if c<len(degree):
                    if check=='year' and len(years)==len(degree):
                        if c==len(degree):
                            s= sentence.upper().index(years[i])
                            e= s+200
                            tex=sentence[s:e]
                        else:
                            s=sentence.upper().index(years[i])
                            e=sentence.upper().index(years[i+1])
                            tex=sentence[s:e]
                            
                            
                    elif check=='degree' or (check=='year' and len(years)!=len(degree)):
                        if c==len(degree):
                            s = sentence.upper().index(degree[i])
                            e = s+200
                            tex=sentence[s:e]
                        else:
                            s = sentence.upper().index(degree[i])
                            e = sentence.upper().index(degree[i+1])                    
                            tex = sentence[s:e] 
                            
                    elif check=='institute':
                        if c==len(degree):
                            if len(degree)>1:
                                s = sentence.upper().index(degree[i-1])
                                e = sentence.upper().index(degree[i])
                                tex=sentence[s:e]
                                
                            if len(degree)==1:
                                s=sentence.upper().index(degree[i])-30
                                e=s+250
                                tex=sentence[s:e]
                        else:
                            if i==0:
                                s=sentence.upper().index(degree[i])
                                e=sentence.upper().index(degree[i+1])
                                tex=sentence[:s]
                                
                            if i>=1:
                                s=sentence.upper().index(degree[i-1])
                                e=sentence.upper().index(degree[i])
                                tex=sentence[s:e]
                                
                if tex!='':
                    d=Education1.year_and_institute(text,qual,degree,institute_names,i,s,e,tex)
                    if len(years)==len(degree):
                        d['YEARS OF PASSING']=years[i]
                else:
                    d=0
                c+=1
                if d!=0:
                    qualifn.append(d)
                    
            if len(match)==0 and len(qualifications)==0:
                d={}
                d['LEVEL']=''
                d['QUALIFICATION_LEVEL']=''
                d['YEAR OF PASSING']=''            
                d['INSTITUTE_NAME']=''
                d['BOARD']=''   
                d['GRADE']=''  
                d['TEXT']=''         
                qualifn.append(d)
            return qualifn
        except:
            logger.exception('erroe in the education qual function')
            traceback.print_exc()
            return []
    
    # this function is used to find the instiute name 
    def institute_name_finder(text,institute_name,degree):
        """ This function finds the institute name in the resume 

        Args:
            text (str): contains the text in which we need to find the institute
            institute_name (set): Contains the list of all the institutes from the database 
            degree (list): It is the list of degree which we found in the resume 

        Returns:          
            list : A list which contains all the institues which we found in the resume  
        """        
                               
        institute_name=[re.sub(r"[,;@#?!&$/%^\s.]+ *","", x.lower())  for x in institute_name]        
        match=[]
        text = re.sub(r"[,;@#?!&$/%^.°\"]+ *"," ", text)
        institute_text=[]
        institute_text_1=[]
        for i in range(1,6):
            s = ngrams(text.split(), i) 
            for j in s:
                institute_text.append(' '.join(j))
        institute_text=[x for x in institute_text if x not in stop_words]     
        
        for y in institute_text:
            y_before=y
            y=re.sub(r"[,;@#?!&$/%^\s.]+ *","", y)
            y=y.lower()
            if y in institute_name:
                # print('found the institue ............',y_before)
                match.append(y_before)
                        
        if len(match)>0:
            return match
        elif len(match)==0:
            degree=[x.lower() for x in degree]
            text=text.lower()
            words_need_to_be_removed=['examination','board','year passed','percentage of marks','cpi','education','year of passing','board','cgpa','percentage','post graduation','under graduation','other qualifications','certifications','programs','pg','ug','description','required criteria','possessed by the candidat','whether eligible or Not','technical consultation training','technical training', 'workshop','technical conference' ,'publication','degree','diploma','specialization','college','year','degree','diploma''qualification','name of college', 'university','degree obtained','discipline','percentage','duration of course','standard','name of institution','years attended','degrees obtained','course','class','educational institutions','profession','degrRAININGee','certificates','diploma','% of marks','cgpa' ,'%','Professional certification','traning','workshop', 'seminar','webinar']
            words_need_to_be_removed.extend(degree)
            text_to_be_searched=text.split()
            for x in words_need_to_be_removed:
                text=text.replace(x,'')
            # print('the text after cleaning is',text)
            text=re.sub(r'\d','',text)
            if len(text)<70:
                try:
                    options = Options()
                    # print('text is empty')
                    options.add_argument("--headless")
                    driverpath = str(BASE_DIR)
                    logpath = str(BASE_DIR) + '/geckodriver.log'  #executable_path=str(BASE_DIR)+'/.wdm/drivers/geckodriver/linux64/v0.33.0/geckodriver'
                    # print('website created start')
                    driver= webdriver.Firefox(options=options, service=FirefoxService(executable_path=str(BASE_DIR)+'/.wdm/drivers/geckodriver/linux64/v0.33.0/geckodriver',log_path=path.devnull))
                    driver.get('https://www.google.com/')
                    # print('website created')
                except:
                    traceback.print_exc()
                try:
                    WebDriverWait(driver,30).until(EC.presence_of_element_located((By.XPATH,'//textarea[@id="APjFqb"]')))
                    driver.find_element('xpath','//textarea[@id="APjFqb"]').send_keys(text)
                    driver.find_element('xpath','//div[@class="o3j99 LLD4me yr19Zb LS8OJ"]').click()
                    driver.find_elements('xpath','//input[@value="Google Search"]')[1].click()
                    search=driver.find_element('xpath','//span[@class="VuuXrf"]').text
                    # print('the match found on google search is',search)
                    driver.quit()
                    if search!=None:
                            match.append(search)
                            return match
                except:
                    traceback.print_exc()
                    driver.quit()
                    # print('some error occured while searching the text on google')
            
        elif len(match)==0:
            institute_text=' '.join([x for x in institute_text])
            match.append(Education1.univer_name(institute_text))   
            return match   
        
          
    
    # this function is used to find the university name.................................
    def univer_name(text):  
        """This function finds the university name in the resume 

        Args:
            text (str): The text of the resume 

        Returns:
            str: The name of the uinversity 
        """        
        check = ['university','Secondary','education','vishwavidyalaya','council','institute','indian','iit','nit','national','technology','science','board','institution']
        t = text.lower().replace(',','').replace('.', ' ').replace('/', ' ').split()
        index = 0
        t_1=''
        t_2=''
        t_3=''
        if any(x in t for x in check)==True:        
            try:          
                for i in check:
                    if i in t:
                        # print('match univeristy', i )
                        index+=t.index(i) 
                        t_2+=i
                        break
                if len(re.findall(r'[A-z]',t[index-1]))==len(t[index-1]):
                    try:
                        if len(re.findall(r'[A-z]',t[index-2]))==len(t[index-2]):
                            t_1+=' '.join(t[index-2:index])
                    except:
                        t_1+=t[index-1]        
                if t[index+1]=='of':
                    if len(re.findall(r'[A-z]',t[index+2]))==len(t[index+2]):
                        if t[index+3]=='and':
                            if len(re.findall(r'[A-z]',t[index+5]))==len(t[index+5]):
                                t_3+=' '.join(t[index+1:index+5])
                        else:
                            t_3+=' '.join(t[index+1:index+3])
            except:
                pass
        return t_1+' '+t_2+' '+t_3
