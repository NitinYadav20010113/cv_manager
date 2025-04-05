import re
import nltk
tokenizer = nltk.data.load('tokenizers/punkt/english.pickle')
from docx import Document
import traceback
#code for getting the project details..............................................................

import spacy


# get project text 
# project heading synonyms
project_headiongs=['RELEVANT ASSIGNMENTS','Employment record relevant to the assignment','Works Undertaken','KEY PROJECTS','Some of the major projects done and activities performed','Major International Projects Undertaken','Major Projects Undertaken','List of projects on which the Personnel has worked','DETAILED DESIGN PROJECT','CO-ORDINATION PROJECT','PREBID DESIGN PROJECT','PROJECTS UNDERTAKEN','SPECIAL PROJECTS','Relevant Project Experience','project'
,'Project Undertaken','Projects','EMPLOYMENT RECORD','List of Projects','NOTEWORTHY ASSIGNMENTS'
,'Employment Record'
,'Employment Record Relevant to the Assignment' 
,'Adequacy for the Assignment'
,'PROFESSIONAL EXPERIENCE'
,'CAREER CONTOUR'
,'PROFESSIONAL ABRIDGMENT'
,'Professional experience record (projects)'
,'EMPLOYMENT RECORD'
,'Projects Accomplished'
,'Major Projects Undertaken'
,'Projects Handled'
,'EMPLOYMENT RECORD RELEVANT TO THE PROJECT'
,'List of projects and Assignments'
,'Occupational Contour'
,'ORGANISATIONAL EXPERIENCE'
,'EMPLOYEMET RECORD'
,'Working Exposure'
,'Key Project Experience'
,'Work Profile'
,'PROJECT WISE - EXPERIENCE'
,'DETAILS OF WORK'
,'Portfolio'
,'WORK HISTORY'
,'LIST OF PROJECTS DESIGNED AND EXECUTED'
,'SCAN OF EMPLOYMENT HISTORY'
,'list of selected projects and my role','name of service or project','name of assignment or project','name and location of project','name of service or project','name of project','name of project','name of assignment','project details','project name','projects','project','title']


projects_list=['\nkey projects\n','\nprojects\n','\nlist of projects\n','\nacademic projects\n','\npersonal projects\n']
end_index_list=['personal details','professional certifications','it skills','training','experience(essential/desirable)','employment record','professional memberships','professional associations','other details','technical  training / workshop','total experience','key qualification','employment record relevant to the assignment','skills and experiences','it forte','accolades','industrial training','other trainings','language skills','areas of expertise','salary details','extra curricular activities','organizational exposure','total period of experience','technical skills','skills','personal details','software proficiency','employment record relevant to the project','software skills','membership of professional societies','current assignment','work experience','declaration','interests and awards','additional qualification','personal vitea','professional history','job profile summary','experience profile','computer skills','project training','total work experience','computer proficiency','certification','key skills','project wise - experience','personal profile','professional qualification','software education','experience','hobbies','achievement','technical proficiency','internship training','affiliations','specialised training','career experience','personal minutiae','professional certification','tranning','courses','academic projects','work history','college projects','computer knowledge','job profile','ph. d thesis title','personal dossier','training received','project profile','visiting fellowship','professional affiliation','areas of expertise','membership in professional associations','general information','softwares and other packages','professional trainings attended','key result areas','career objective','job responsibilities','extra carricular activies','inplant training attended','projects during college','extra course','areas of interest','skills and abilities','skills & proficiencies','technical details','industrial experience','work experience/ professional career','work responsibilities','work experience & responsibilities in hand','technicalskills','final internship','software','carrier vision','memberships','software exposure','languages and interest','extra qualification','professional chronicle','academic achievements','training completed','technical knowledge purview','personal qualities','professional experience','achievements','it skills','education summary','personal profile','personal details','personal info','education','training and development','industrial training','others','academics','strengths','it forte','awards, participations, certificates and achievements','training program','m.tech dissertation','internship trainings','computer/technical skills','professional certifications','internship','training and development','company profile','achivement','extra-curricular activities','professional enhancements','technical skills','job experience','languages','modelling skills','summer training','academia','programmes / conferences / workshops attended','professional experience summary','software skills','interpersonal skills','membership of professional bodies','internships','work experience','certificate / experience','educational qualification','co-curricular activities','computer skills','certifications']

client_list=['client','client worked with','client name']

position_list=['positions held','position held','position']
# project details regex
project_name_regex=r'(list of selected projects and my role|name of service or project|name of assignment or project|name and location of project|name of service or project|name of project|name of project|name of assignment|project details|project name|project|title|project detail)'
# client regex
client_name_regex=r'(client|client worked with|client name|clients)'
# designation regex
position_name_regex=r'(positions held|position held|position)'


# getting employer.............................................

def getemployer(text):
    employer=[]
    try:
        for emat in re.finditer(r'(Employer|employer|EMPLOYER)'+'(\n+|\t+|\s)*(:|-|:-){1,}',text.lower()):
            s_index=emat.end()
            employer_list=[x.strip() for x in tokenizer.tokenize(text[s_index:s_index+100])[0].split('\n')[0:1] if x!='']
            if len(employer_list)!=0:
                employer.append(employer_list[0])
    except:
        pass
    return employer

# getting client...............................................
def getclient(text):
    client=[]
    try:
        for cmat in re.finditer(client_name_regex+'(\n+|\t+|\s)*(:|-|:-){1,}',text.lower()):
            s_index=cmat.end()
            client_list=[x.strip() for x in tokenizer.tokenize(text[s_index:s_index+100])[0].split('\n') if x!='']
            if len(client_list)!=0:
                client.append(client_list[0])
    except:
        pass
    return client


#getting position...............................................
def getposition(text):
    position=[]
    try:
        for pmat in re.finditer(position_name_regex+'(\n+|\t+|\s)*(:|-|:-){1,}',text.lower()):
            s_index=pmat.end()
            position_list=[x.strip() for x in tokenizer.tokenize(text[s_index:s_index+100])[0].split('\n') if x!='']
            if len(position_list)!=0:
                position.append(position_list[0])
    except:
        pass
    return position


#getting all project details........................................
def getprojectdetails(text):
    # # print(text)
    text=text.replace('  ',' ')
    project={}
    p_list=[]
    c_list=[]
    e_list=[]
    pos_list=[]
    projects=[]
    end_index=len(text)
    nlp=spacy.load('en_core_web_lg')
    for x in projects_list:
        
        if x in text.lower():
            s_index=text.lower().index(x)
            for y in end_index_list:
                p_text=text[s_index:]
                ma=re.search(r'\n{}(\n|:)'.format(y),p_text.lower())
                if ma!=None:
                    new_index=s_index+ma.start()
                    if new_index<=end_index:
                        end_index=s_index+ma.start()
            
            try:
                project_text=text[s_index:end_index]
                projects=re.split(r'(?<!\d)\d{1,2}(?!\d)\.\s',project_text)
                if len(projects)>1:    
                    project['title']=projects
                    project['client']=[]
                    project['employer']=[]
                    project['position']=[]
                    return project
                
                elif len(projects)<2:
                    projects=re.split(r'\ne',project_text)
                    if len(projects)>1:
                        project['title']=projects
                        project['client']=[]
                        project['employer']=[]
                        project['position']=[]
                        return project
                    
                    else:
                        doc=nlp(project_text)
                        l=[]
                        for sent in doc.sents:
                            l.append(sent)
                        if len(l)>2:
                            project['title']=l
                            project['client']=[]
                            project['employer']=[]
                            project['position']=[]
                            return project
            except:
                break
    
   
    
    try:
        # getting project heading
        for x in project_headiongs:
            if x.lower() in text.lower():
                # # print('x======',x)
                s_index=text.lower().index(x.lower())
                p_text=text[s_index:]
                # getting project name
                for pnmat in re.finditer(project_name_regex+'(\n+|\t+|\s|\d)*(:|-|:-){1,}',p_text.lower()):
                    
                    s_index=pnmat.end()
                    project_list=[x.strip() for x in tokenizer.tokenize(p_text[s_index:])[0].split('\n')[0:3] if x!='']
                    if len(project_list)!=0:
                        p_list.append(project_list[0])
                e_list.append(getemployer(p_text))
                c_list.append(getclient(p_text))
                pos_list.append(getposition(p_text))
                break
        if len(p_list)==0:
            p_list=[]
            c_list=[]
            e_list=[]
            pos_list=[]
            for pnmat in re.finditer(project_name_regex+'(\n+|\t+|\s|\d)*(:|-|:-){1,}',text.lower()):
                s_index=pnmat.end()
                project_list=[x.strip() for x in tokenizer.tokenize(text[s_index:])[0].split('\n')[0:3] if x!='']
                if len(project_list)!=0:
                    p_list.append(project_list[0])
            e_list.append(getemployer(text))
            c_list.append(getclient(text))
            pos_list.append(getposition(text))
        
        # print('the length of p_list is',len(p_list))
        # print('the length of client list is',len(c_list[0]))
        # print('the length of postion list is',len(pos_list[0]))
        # print('the length of e list is ',len(e_list[0]))
        if len(p_list)!=0:
            project['title']=p_list
            if len(c_list)!=0:
                project['client']=c_list[0]
            else:
                project['client']=['']*len(p_list)

            if len(e_list)!=0:
                project['employer']=e_list[0]
            else:
                project['employer']=['']*len(p_list)
            if len(pos_list)!=0:
                project['position']=pos_list[0]
            else:
                project['position']=['']*len(p_list)
            return project
        else:
            for x in projects_list:
                l_text=text.lower()
                if x in l_text:
                    s_index=l_text.index(x)
                    p_text=l_text[s_index:]
                    projects=re.split(r'(?<!\d)\d{1,2}(?!\d)\.\s',p_text)
                    if len(projects)>1:    
                        projects[-1]=projects[-1][0:20]
                        project['title']=projects
                        project['client']=[]
                        project['employer']=[]
                        project['position']=[]
                        return project
                    
                    elif len(projects)<2:
                        l_text=text.lower()
                        s_index=l_text.index(x)
                        p_text=l_text[s_index:]
                        projects=re.split(r'\ne',p_text)
                        if len(projects)>1:
                            projects[-1]=projects[-1][0:20]
                            project['title']=projects
                            project['client']=[]
                            project['employer']=[]
                            project['position']=[]
                            return project
                        
                        else:
                            projects[-1]=projects[-1][0:20]
                            project['title']=[]
                            project['client']=[]
                            project['employer']=[]
                            project['position']=[]
                            return project
                        
                    else:
                        projects[-1]=projects[-1][0:20]
                        project['title']=[]
                        project['client']=[]
                        project['employer']=[]
                        project['position']=[]
                        return project
                    
                    
                else:
                    projects[-1]=projects[-1][0:20]
                    project['title']=[]
                    project['client']=[]
                    project['employer']=[]
                    project['position']=[]
                    return project
    except:
        project['title']=[]
        project['client']=[]
        project['employer']=[]
        project['position']=[]
        return project
    # print('the project detail is',project)
    # print('the project detail is',project)
    # print('the project detail is ',project)
    






    
# def table_no_finder(path='/home/nitinyadav/Downloads/cvdata/Aniket Shama.docx'):
   
#     main_list=[project_list,client_list,position_list]
    
    
#     document = Document(path)
#     table = document.tables

#     for x in range(len(table)):
#         table =document.tables[x]
        
#         data = []

#         keys = None
#         for i, row in enumerate(table.rows):
#             text = (cell.text for cell in row.cells)

#             if i == 0:
#                 keys = tuple(text)
#                 continue
#             row_data = dict(zip(keys, text))
#             data.append(row_data)
#         df = pd.DataFrame(data)
#         list_of_column=list(df.columns)

#         count=0
#         for y in list_of_column:
#             for z in list_of_all_headings:
#                 match=re.search(z.lower(),y.lower())
#                 if match!=None:
#                     break
#             if match!=None:
#                 count+=1
#         if count>=2:
#             break

#     data=table_data_extractor(document,x,textt)    
#     return  data





# def table_data_extractor(document,table_no):
#     dic={}
#     qualification=[]
#     collage=[]
#     university=[]
#     year_of_passing=[]
#     percentage=[]
#     table = document.tables[table_no]
#     data = []
#     main=[]

#     keys = None
#     for i, row in enumerate(table.rows):
#         text = (cell.text for cell in row.cells)

#         if i == 0:
#             keys = tuple(text)
#             continue
#         row_data = dict(zip(keys, text))
#         data.append(row_data)
#     df = pd.DataFrame(data)
#     list_of_column=list(df.columns)
#     for y in list_of_column:
#         if y.lower() in qualification_list:
#         qualification.extend(df[y].to_list())
#         if y.lower() in college_list:
#             collage.extend(df[y].to_list())
#         if y.lower() in year_of_passing_list:
#             year_of_passing.extend(df[y].to_list())
#         if y.lower() in percentage_list:
#             percentage.extend(df[y].to_list())
#         if y.lower() in university_list:
#             university.extend(df[y].to_list())
            
            
    
#     for x in range(len(qualification)):
#         dic['QUALIFICATION_LEVEL']=qualification[x]
#         dic['YEAR OF PASSING']=year_of_passing[x]
#         dic['INSTITUTE_NAME']=collage[x]
#         dic['BOARD']=university[x]
#         dic['GRADE']=percentage[x]
#         dic['TEXT']=qualification[x]+year_of_passing[x]+collage[x]+university[x]+percentage[x] 
#         main.append(dic)
#         dic={}
        
#     print('the education details is ',main)
    

